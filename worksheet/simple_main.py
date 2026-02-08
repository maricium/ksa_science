#!/usr/bin/env python3
"""
Simplified worksheet generator - generates worksheets for all lessons
"""

import openai
import pandas as pd
from docx import Document
import re
import os
from pathlib import Path

# Configuration
openai.api_key = os.getenv("OPENAI_API_KEY", "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")

def extract_lesson_data(excel_file):
    """Extract lesson data from Excel file"""
    try:
        df = pd.read_excel(excel_file, skiprows=4)  # Skip header rows
        lessons = []
        
        for _, row in df.iterrows():
            if pd.notna(row.iloc[1]) and str(row.iloc[1]).strip():  # Check lesson title
                lesson = {
                    'code': str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else '',
                    'title': str(row.iloc[1]).strip(),
                    'knowledge': str(row.iloc[2]).strip() if pd.notna(row.iloc[2]) else '',
                    'skills': str(row.iloc[4]).strip() if pd.notna(row.iloc[4]) else ''
                }
                lessons.append(lesson)
        
        return lessons
    except Exception as e:
        print(f"Error reading {excel_file}: {e}")
        return []

def generate_worksheet(lesson, subject="Chemistry"):
    """Generate a single worksheet using AI"""
    prompt = f"""Create a GCSE {subject} worksheet with 21 questions (15 minutes) on: {lesson['title']}

Learning Objectives:
- Knowledge: {lesson['knowledge']}
- Skills: {lesson['skills']}

Requirements:
- Questions 1-7: Low Demand (1-2 marks)
- Questions 8-14: Standard Demand (2-3 marks)  
- Questions 15-21: High Demand (3-4 marks)

Format: Questions first, then complete mark scheme with model answers for ALL 21 questions."""

    try:
        client = openai.OpenAI(api_key=openai.api_key)
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=6000
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error generating worksheet: {e}")
        return None

def create_word_document(content, lesson):
    """Create Word document from AI content"""
    doc = Document()
    
    # Add title
    doc.add_heading(lesson['title'], 0)
    doc.add_paragraph("")
    
    # Process content
    lines = content.split('\n')
    in_questions = True
    question_pattern = re.compile(r'^\d+\.')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for mark scheme section
        if "mark scheme" in line.lower() or "marking" in line.lower():
            in_questions = False
            doc.add_page_break()
            doc.add_heading("MARK SCHEME", 1)
            continue
        
        if in_questions and question_pattern.match(line):
            # Add question
            doc.add_paragraph(line)
            
            # Add answer lines
            question_num = int(line.split('.')[0])
            num_lines = 2 if question_num <= 7 else 3 if question_num <= 14 else 4
            
            for _ in range(num_lines):
                p = doc.add_paragraph()
                p.add_run("_" * 120)
            
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    
    return doc

def main():
    """Main function to generate all worksheets"""
    print("🔄 Generating worksheets...")
    
    # Excel files to process
    excel_files = [
        "../Links/C4.1 All Resources/c4.1 unit overview xslx/C4.1 Unit Overview (4).xlsx",
        "../Links/C4.2 All Resources/C4.2 Unit Overview (1).xlsx",
        "../Links/C4.3 All Resources/C4.3 Unit Overview (1).xlsx",
        "../Links/C4.4 All Resources/C4.4 Unit Overview.xlsx"
    ]
    
    # Create worksheets directory
    worksheets_dir = Path("worksheets")
    worksheets_dir.mkdir(exist_ok=True)
    
    total_lessons = 0
    for excel_file in excel_files:
        if os.path.exists(excel_file):
            lessons = extract_lesson_data(excel_file)
            print(f"📊 {excel_file}: Found {len(lessons)} lessons")
            
            for lesson in lessons:
                print(f"Creating: {lesson['title']}")
                
                # Generate worksheet content
                content = generate_worksheet(lesson)
                if not content:
                    continue
                
                # Create Word document
                doc = create_word_document(content, lesson)
                
                # Save with unit organization
                unit_code = lesson['code'].split('.')[0] + '.' + lesson['code'].split('.')[1] if '.' in lesson['code'] else 'Unknown'
                unit_dir = worksheets_dir / unit_code
                unit_dir.mkdir(exist_ok=True)
                
                filename = f"{lesson['code']}_{lesson['title'].replace(' ', '_')}_worksheet.docx"
                doc.save(unit_dir / filename)
                print(f"✅ Created: {unit_code}/{filename}")
                total_lessons += 1
    
    print(f"\n🎉 Generated {total_lessons} worksheets successfully!")

if __name__ == "__main__":
    main()

