#!/usr/bin/env python3
"""
Generate mini whiteboard question documents
10 questions per lesson: fill-in-the-blank, multiple choice, etc.
"""

import os
import sys
import openai
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from main import extract_lesson_data, extract_objectives_from_pptx

# Configuration - Change this to generate questions for different units
UNIT_CODE = "C4.2"  # e.g., "B3.2", "C4.2", "C4.1", etc.
TARGET_LESSON = "C4.2.5"  # Set to None to generate for all lessons in the unit

# Configuration
openai.api_key = os.getenv("OPENAI_API_KEY", "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")

def generate_mini_whiteboard_questions(lesson, pptx_objectives=None):
    """Generate 10 mini whiteboard questions using OpenAI"""
    
    # Build context
    context = f"""
LESSON INFORMATION:
Title: {lesson['title']}

Knowledge students need to know:
{lesson.get('knowledge', '')}

Skills students need to demonstrate:
{lesson.get('skills', '')}
"""
    
    if pptx_objectives:
        context += f"\n\nLearning objectives from the lesson:\n"
        context += "\n".join([f"- {obj}" for obj in pptx_objectives])
    
    prompt = f"""{context}

CREATE 10 MINI WHITEBOARD QUESTIONS FOR THIS LESSON

REQUIREMENTS:
1. Create exactly 10 questions
2. Mix of question types:
   - Fill-in-the-blank (e.g., "The process of _____ is when...")
   - Multiple choice (e.g., "What is X? A) option1 B) option2 C) option3 D) option4")
   - True/False (e.g., "True or False: ...")
   - One-word answers (e.g., "What is the name of...?")
   - Short definitions (e.g., "Define the term 'X'")
   - Match the words to definitions (e.g., "Match the words to their definitions: Word1 - Definition1, Word2 - Definition2, etc." Students draw lines to connect them)
3. Questions should be QUICK to answer (30 seconds or less)
4. Questions should test KEY facts and concepts from the lesson
5. Questions should be based on the knowledge and skills provided above
6. Use simple, clear language suitable for GCSE students
7. Include at least ONE "match the words to definitions" question in the set

FORMAT:
For each question, provide:
- Question number (1-10)
- Question text
- Answer (for fill-in-the-blank, provide the missing word/phrase; for multiple choice, provide the correct letter; for true/false, provide True or False)

OUTPUT FORMAT:
1. [Question type]: [Question text]
   Answer: [Answer]

2. [Question type]: [Question text]
   Answer: [Answer]

... (continue for all 10 questions)

Make sure questions are:
- Quick to read and answer
- Test important facts from the lesson
- Suitable for students to write on mini whiteboards
- Clear and unambiguous
"""

    try:
        client = openai.OpenAI(api_key=openai.api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error generating questions: {e}"

def create_mini_whiteboard_document(questions_text, lesson):
    """Create Word document with questions and answers"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading(f"Mini Whiteboard Questions: {lesson['title']}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add lesson code
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(f"Lesson: {lesson['code']}")
    code_run.bold = True
    code_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Parse questions and answers
    import re
    questions = []
    answers = []
    
    # Split by question numbers
    question_blocks = re.split(r'(\d+\.)', questions_text)
    
    for i in range(1, len(question_blocks), 2):
        if i + 1 < len(question_blocks):
            question_num = question_blocks[i].strip()
            content = question_blocks[i + 1].strip()
            
            # Split question and answer
            if 'Answer:' in content or 'answer:' in content:
                parts = re.split(r'[Aa]nswer:\s*', content, 1)
                question_text = parts[0].strip()
                answer_text = parts[1].strip() if len(parts) > 1 else ""
            else:
                question_text = content
                answer_text = ""
            
            # Clean up question text (remove question type prefix if present)
            question_text = re.sub(r'^\[.*?\]:\s*', '', question_text)
            
            if question_text:
                questions.append((question_num, question_text))
                answers.append((question_num, answer_text))
    
    # Add questions
    for q_num, q_text in questions:
        # Check if it's a "match the words" question
        if "match" in q_text.lower() and ("word" in q_text.lower() or "definition" in q_text.lower()):
            # Format as a matching question with two columns
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{q_num} {q_text}")
            q_run.font.size = Pt(12)
            q_run.bold = True
            doc.add_paragraph()  # Spacing
            
            # Try to extract words and definitions from the question
            # Format: "Match: Word1 - Definition1, Word2 - Definition2"
            import re
            match_parts = re.findall(r'([A-Za-z\s]+)\s*-\s*([A-Za-z\s,\.]+)', q_text)
            if match_parts:
                # Create a table for matching
                table = doc.add_table(rows=len(match_parts), cols=2)
                table.style = 'Light Grid Accent 1'
                
                for i, (word, definition) in enumerate(match_parts):
                    table.rows[i].cells[0].text = word.strip()
                    table.rows[i].cells[1].text = definition.strip()
                    # Make text smaller for table
                    for cell in table.rows[i].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
        else:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{q_num} {q_text}")
            q_run.font.size = Pt(12)
            q_run.bold = True
        doc.add_paragraph()  # Spacing
    
    # Add page break before answers
    doc.add_page_break()
    
    # Add answers heading
    answers_heading = doc.add_heading("ANSWERS", level=1)
    answers_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Add answers
    for a_num, a_text in answers:
        if a_text:
            a_para = doc.add_paragraph()
            a_run1 = a_para.add_run(f"{a_num} ")
            a_run1.bold = True
            a_run2 = a_para.add_run(a_text)
            a_run2.font.size = Pt(12)
            doc.add_paragraph()  # Spacing
    
    return doc

def generate_for_unit(unit_code, excel_file, lesson_resources_folder, target_lesson=None):
    """Generate mini whiteboard questions for a specific unit"""
    if target_lesson:
        print(f"\n🔄 Generating mini whiteboard questions for {unit_code} (target: {target_lesson})...")
    else:
        print(f"\n🔄 Generating mini whiteboard questions for {unit_code}...")
    
    if not os.path.exists(excel_file):
        print(f"❌ Excel file not found: {excel_file}")
        return 0, 0
    
    # Extract lessons from Excel
    lessons = extract_lesson_data(excel_file)
    print(f"📊 Found {len(lessons)} lessons in {unit_code}")
    
    if not lessons:
        print(f"❌ No lessons found in Excel file for {unit_code}")
        return 0, 0
    
    successful = 0
    failed = 0
    
    for lesson in lessons:
        lesson_code = lesson.get('code', '').strip()
        lesson_title = lesson.get('title', '').strip()
        
        # Handle case where lesson code might be in title column
        unit_prefix = unit_code.split('.')[0] + '.' + unit_code.split('.')[1] if '.' in unit_code else unit_code
        if not lesson_code and lesson_title and lesson_title.startswith(unit_prefix):
            lesson_code = lesson_title
            lesson_title = lesson.get('knowledge', '').strip().split('\n')[0] if lesson.get('knowledge') else lesson_title
        
        if not lesson_code or not lesson_code.startswith(unit_prefix):
            continue
        
        # Filter by target_lesson if specified
        if target_lesson and lesson_code != target_lesson:
            continue
        
        print(f"\n📝 Creating mini whiteboard questions: {lesson_code} - {lesson_title}")
        
        try:
            # Try to extract objectives from PowerPoint
            pptx_objectives = None
            pptx_objectives = extract_objectives_from_pptx(lesson_code, lesson_resources_folder)
            
            # Generate questions
            questions_text = generate_mini_whiteboard_questions(lesson, pptx_objectives)
            
            if not questions_text or "Error" in questions_text:
                print(f"❌ Failed to generate questions for {lesson_code}")
                failed += 1
                continue
            
            # Create Word document
            doc = create_mini_whiteboard_document(questions_text, {
                'code': lesson_code,
                'title': lesson_title
            })
            
            # Save document
            clean_code = lesson_code.replace('/', '_').replace('\\', '_')
            clean_title = lesson_title.replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
            filename = f"{clean_code}_{clean_title}_mini_whiteboard_questions.docx"
            
            doc.save(output_dir / filename)
            print(f"✅ Created: {filename}")
            successful += 1
            
        except Exception as e:
            print(f"❌ Error creating questions for {lesson_code}: {e}")
            import traceback
            traceback.print_exc()
            failed += 1
            continue
    
    return successful, failed

def find_unit_paths(unit_code):
    """Find Excel file and lesson resources folder for a unit code"""
    lesson_resources = Path("../Lesson Resources")
    if not lesson_resources.exists():
        lesson_resources = Path("Lesson Resources")
    
    # Search for Excel file
    excel_files = list(lesson_resources.rglob(f"**/Unit Guidance/**/{unit_code}*.xlsx"))
    if not excel_files:
        excel_files = list(lesson_resources.rglob(f"**/Unit Guidance/**/*{unit_code}*.xlsx"))
    
    excel_file = str(excel_files[0]) if excel_files else None
    
    # Find lesson resources folder (parent of Unit Guidance)
    if excel_file:
        excel_path = Path(excel_file)
        # Go up from Unit Guidance to find Lesson Resources folder
        unit_guidance = excel_path.parent
        lesson_folder = unit_guidance.parent.parent / "Lesson Resources"
        if not lesson_folder.exists():
            # Try alternative structure
            lesson_folder = unit_guidance.parent / "Lesson Resources"
    else:
        lesson_folder = None
    
    return excel_file, str(lesson_folder) if lesson_folder and lesson_folder.exists() else None

def main():
    """Generate mini whiteboard questions for the unit specified in UNIT_CODE constant"""
    global output_dir
    
    # Create output directory
    output_dir = Path("worksheets/pptQuestions")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Find paths for the unit
    excel_file, lesson_folder = find_unit_paths(UNIT_CODE)
    
    if not excel_file:
        print(f"❌ Could not find Excel file for unit {UNIT_CODE}")
        print(f"   Searched in: Lesson Resources/**/Unit Guidance/**/{UNIT_CODE}*.xlsx")
        return
    
    if not lesson_folder:
        print(f"⚠️  Could not find Lesson Resources folder for {UNIT_CODE}")
        print(f"   Will continue without PowerPoint objectives")
        lesson_folder = None
    
    # Generate for the unit
    successful, failed = generate_for_unit(UNIT_CODE, excel_file, lesson_folder, TARGET_LESSON)
    
    print(f"\n{'='*60}")
    print(f"✅ Success: {successful} | ❌ Failed: {failed}")
    print(f"📁 Files saved to: worksheets/pptQuestions/")

if __name__ == "__main__":
    main()

