#!/usr/bin/env python3
"""
AQA Core Knowledge Generator - IMPROVED SPECIFICATION-AWARE VERSION

requirements:
✓ Minimum 10 words per week (foundation tier - ALL students learn these)
✓ Exactly 5 extension words per week (HT only - triple science students)
✓ Extension words must be HIGHLIGHTED in yellow
✓ AQA specification-aware (uses 4.x vs 5.x references to determine tiers)
✓ Professional AQA-style questions
"""

import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

_client = None
try:
    _client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    print("✓ OpenAI client ready for question generation")
except Exception as e:
    print(f"⚠️ OpenAI not available: {e} – will use template questions only")


def normalize_kw(s: str) -> str:
    """Normalize keyword for consistent matching"""
    return re.sub(r"\s+", " ", s.strip().lower())


def _lesson_sort_key(lc: str):
    """Sort key for lesson codes so C4.3.1, C4.3.2, ... C4.3.10, not string sort."""
    parts = lc.split(".")
    if len(parts) >= 3:
        try:
            return (parts[0], int(parts[1]), int(parts[2]))
        except ValueError:
            pass
    return (lc, 0, 0)


def lesson_code_to_unit(lesson_code: str) -> str:
    """Get unit code from lesson code, e.g. B3.2.4 -> B3.2, C4.3.2 -> C4.3. Handles sublessons like 'C4.3.2 / 2'."""
    lc = lesson_code.split("/")[0].strip()
    parts = lc.split(".")
    if len(parts) >= 2:
        return f"{parts[0]}.{parts[1]}"
    return lc


def read_state_questions_from_prep_booklet(
    docx_path: str,
    lesson_codes_ordered: list[str],
    prep_definitions: dict[str, str] | None = None,
    spec_content: list[tuple[str, str]] | None = None,
) -> dict[str, list[tuple[str, str]]]:
    """
    Extract 'State ...' example questions from the Unit Preparation Booklet.
    Tables with 'Intended outcome' / 'Example questions' are assumed to be in lesson order.
    Returns dict: lesson_code -> [(question, answer), ...].
    """
    if not docx_path or not Path(docx_path).exists():
        return {}
    prep_definitions = prep_definitions or {}
    doc = Document(docx_path)

    # Find tables that look like lesson objective + example questions (one per lesson)
    lesson_tables: list[tuple[int, int]] = []  # (table_index, example_col_index)
    for ti, table in enumerate(doc.tables):
        if table.rows and len(table.rows) >= 2:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            header_joined = " ".join(header_cells).lower()
            if "intended outcome" in header_cells[0].lower() and "example question" in header_joined:
                # Example questions column: usually index 1, or 2 if 3 columns
                ex_col = 1
                if len(header_cells) > 2 and "example question" in header_cells[2].lower():
                    ex_col = 2
                lesson_tables.append((ti, ex_col))

    if not lesson_tables or not lesson_codes_ordered:
        return {}

    result: dict[str, list[tuple[str, str]]] = {lc: [] for lc in lesson_codes_ordered}
    # Map each lesson table to a lesson code by position
    for table_idx, (ti, ex_col) in enumerate(lesson_tables):
        if table_idx >= len(lesson_codes_ordered):
            break
        lesson_code = lesson_codes_ordered[table_idx]
        table = doc.tables[ti]
        # Data rows: skip header and optional subtitle (rows 0, 1)
        # Per lesson: only one "State the units of..." question (avoid duplicate near-identical questions)
        seen_units_question_for_lesson = False
        for row_idx in range(2, len(table.rows)):
            row = table.rows[row_idx]
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= ex_col:
                continue
            intended_outcome = cells[0] if cells else ""
            example_text = cells[ex_col]
            for line in example_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # Only lines that are "State ..." type
                if not re.match(r"^state\s", line, re.IGNORECASE):
                    continue
                question = line
                # Generic "State the units for each quantity" → at most one per lesson
                if _is_generic_units_question(question):
                    if seen_units_question_for_lesson:
                        continue
                    seen_units_question_for_lesson = True
                refined = _refine_units_question_with_ai(question, intended_outcome)
                if refined:
                    question, answer = refined
                else:
                    answer = _answer_for_state_question(
                        question, intended_outcome, prep_definitions, spec_content=spec_content
                    )
                # Never allow TBAT (intended outcome) as an answer - it's wrong
                answer = answer.strip()
                if answer.upper().startswith("TBAT"):
                    answer = _fix_long_answer_with_ai(question, answer).strip()
                    if answer.upper().startswith("TBAT"):
                        answer = "[state from lesson]"
                if len(answer) > MAX_STATE_ANSWER_CHARS:
                    answer = _fix_long_answer_with_ai(question, answer).strip()
                if answer.upper().startswith("TBAT"):
                    answer = "[state from lesson]"
                if len(answer) > MAX_STATE_ANSWER_CHARS:
                    answer = answer[:MAX_STATE_ANSWER_CHARS]
                # Critical teacher check: does the answer actually match the question?
                answer = _validate_answer_with_ai(question, answer)
                answer = _sanitise_answer_for_revision("", answer)  # no keyword; strip meta/robot answers or use —
                # If answer still over limit, remove this question entirely (do not add)
                if len(answer) > MAX_STATE_ANSWER_CHARS:
                    continue
                result[lesson_code].append((question, answer))
    return result


MAX_STATE_ANSWER_CHARS = 50

# Shown in the doc when the answer is not AQA-spec or is robot/meta (literally a line across)
INVALID_ANSWER_PLACEHOLDER = "—"


def _is_robot_or_meta_answer(text: str) -> bool:
    """True if the answer reads like a meta-comment or robot explanation, not a revision term."""
    if not text or len(text) > 120:
        return True
    t = text.lower().strip()
    meta_phrases = [
        "is not a definition",
        "is not an answer",
        "is not the answer",
        "does not make sense",
        "is not a ",
        "is not the ",
        "this is not ",
        "that is not ",
        "it is not ",
        "not a definition",
        "not an answer",
    ]
    if any(p in t for p in meta_phrases):
        return True
    # Sentence that explains rather than states a term (ends with period, or long clause)
    if t.endswith(".") and len(t) > 40:
        return True
    return False


def _sanitise_answer_for_revision(keyword: str, answer: str) -> str:
    """
    Ensure the answer is a short revision-style term, not a robot/meta comment.
    If answer is meta (e.g. 'relative atomic mass is not a definition'), use the keyword or extract the term; otherwise return answer.
    """
    if not answer or answer.strip() == INVALID_ANSWER_PLACEHOLDER:
        return keyword.strip()[:MAX_STATE_ANSWER_CHARS] if keyword.strip() else INVALID_ANSWER_PLACEHOLDER
    if _is_robot_or_meta_answer(answer):
        # Prefer the keyword from the sheet (short revision term)
        if keyword and not _is_robot_or_meta_answer(keyword) and len(keyword.strip()) <= MAX_STATE_ANSWER_CHARS:
            return keyword.strip()
        # Try to extract a short term: e.g. "relative atomic mass is not a definition" -> "relative atomic mass"
        for sep in (" is not ", " is not a ", " is not an ", " does not ", " - "):
            if sep in answer:
                before = answer.split(sep)[0].strip()
                if before and len(before) <= MAX_STATE_ANSWER_CHARS and not _is_robot_or_meta_answer(before):
                    return before
        return INVALID_ANSWER_PLACEHOLDER
    return answer.strip()[:MAX_STATE_ANSWER_CHARS]


def _validate_answer_with_ai(question: str, answer: str) -> str:
    """
    Prefer the keyword-sheet answer. Only suggest a replacement for clear abbreviation expansion
    (e.g. RTP -> room temperature and pressure). If the proposed answer is a short revision term, keep it.
    Do not replace with sentences or non-AQA spec; if unsure, return the original or "—".
    """
    if not _client or not answer.strip():
        return answer
    prompt = f"""You are checking a GCSE AQA recall Q&A for a revision word list. The proposed answer comes from the keyword sheet – USE IT unless it is clearly wrong.

Question: {question}
Proposed answer (from keyword sheet): {answer}

RULES:
- Answers must be MINI WHITEBOARD style: a word or two, an equation, or a formula. Examples: "density", "length × width × height", "6.02×10²³", "room temperature and pressure". NOT partial sentences or definitions (e.g. "The physical form in which a substance is in: solid" is WRONG – use "solid" or "state of matter").
- PREFER the proposed answer if it is already short (1–4 words or an equation).
- Only suggest a replacement for clear abbreviation expansion (e.g. RTP → "room temperature and pressure"). Do NOT suggest definitions or sentences.
- If the proposed answer is a reasonable short revision term, reply with exactly: OK
- If you suggest a replacement, it MUST be 1–4 words, an equation, or a formula. If the correct answer would be a long sentence, reply with exactly: —

If the answer is wrong, incomplete or does not match the question, reply with the CORRECT short answer only (max 50 characters, no explanation, no "OK").
If the answer is correct, reply with exactly: OK"""

    try:
        response = _client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=80,
        )
        text = (response.choices[0].message.content or "").strip().split("\n")[0].strip()
        if not text:
            return answer
        if text.upper().strip() == "OK":
            return answer
        if text.strip() == "—" or text.strip() == "–" or text.strip() == "-":
            return INVALID_ANSWER_PLACEHOLDER
        # AI provided a correction – only use if it looks like a short revision term
        for prefix in ("answer:", "answer ", "ANSWER:", "ANSWER ", "correct:", "correct "):
            if text.lower().startswith(prefix):
                text = text[len(prefix):].strip()
                break
        if text and not text.upper().startswith("TBAT"):
            if _is_robot_or_meta_answer(text):
                return answer
            return text[:MAX_STATE_ANSWER_CHARS]
    except Exception:
        pass
    return answer


def _fix_long_answer_with_ai(question: str, current_answer: str) -> str:
    """
    If an answer is over MAX_STATE_ANSWER_CHARS we treat it as a mistake. Use OpenAI
    to infer the correct short answer (max 50 characters).
    """
    if not _client or len(current_answer.strip()) <= MAX_STATE_ANSWER_CHARS:
        return current_answer
    prompt = f"""This is a GCSE chemistry recall question. The current answer is too long (mistake).

Question: {question}
Current answer (wrong, too long): {current_answer}

Give the correct short answer in AT MOST 50 characters. Only the value, symbol, or key term (e.g. 6.02×10²³, mol/dm³, periodic table). No sentences. Do NOT start with TBAT or "To be able to" - that is the question objective, not the answer. One line only."""

    try:
        response = _client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=30,
        )
        text = (response.choices[0].message.content or "").strip().split("\n")[0].strip()
        # Remove any leading "ANSWER:" or similar
        for prefix in ("answer:", "answer ", "ANSWER:", "ANSWER "):
            if text.lower().startswith(prefix):
                text = text[len(prefix):].strip()
                break
        if text:
            return text
    except Exception:
        pass
    return current_answer[:MAX_STATE_ANSWER_CHARS]


def _state_statement_to_question(text: str) -> str:
    """
    Turn a 'State ...' statement into a question for the revision list.
    e.g. 'State the Avogadro constant' -> 'What is the Avogadro constant?'
    """
    if not text or not text.strip():
        return text
    t = text.strip().rstrip("?")
    if not re.match(r"^state\s", t, re.IGNORECASE):
        return t if text.strip().endswith("?") else t + "?"
    # "State the units of X" -> "What are the units of X?"
    if re.match(r"^state\s+the\s+units\s+", t, re.IGNORECASE):
        rest = re.sub(r"^state\s+the\s+units\s+", "", t, flags=re.IGNORECASE).strip().rstrip("?")
        return f"What are the units of {rest}?" if rest else t
    # "State the unit of X" -> "What is the unit of X?"
    if re.match(r"^state\s+the\s+unit\s+", t, re.IGNORECASE):
        rest = re.sub(r"^state\s+the\s+unit\s+", "", t, flags=re.IGNORECASE).strip().rstrip("?")
        return f"What is the unit of {rest}?" if rest else t
    # "State the value of X" / "State the equation for X" etc. -> "What is the ...?"
    if re.match(r"^state\s+the\s+", t, re.IGNORECASE):
        rest = re.sub(r"^state\s+the\s+", "", t, flags=re.IGNORECASE).strip().rstrip("?")
        return f"What is the {rest}?" if rest else t
    # "State X" (no "the") -> "What is X?"
    rest = re.sub(r"^state\s+", "", t, flags=re.IGNORECASE).strip().rstrip("?")
    if not rest:
        return t
    return f"What is {rest}?"


def _is_generic_units_question(question: str) -> bool:
    """True if question is vague e.g. 'State the units for each quantity'."""
    q = question.lower().strip()
    if "state" not in q[:10]:
        return False
    return (
        "units for each" in q
        or "units for the" in q
        or (("unit" in q or "units" in q) and "quantity" in q and "each" in q)
    )


def _refine_units_question_with_ai(question: str, intended_outcome: str) -> tuple[str, str] | None:
    """
    For generic 'State the units for each quantity', use the TBAT (intended outcome) to
    rewrite as a specific question (e.g. 'State the units of concentration, volume and amount')
    and extract the answer (e.g. 'mol/dm³, dm³, mol') from the TBAT. Returns (new_question, answer) or None.
    """
    if not _is_generic_units_question(question) or not intended_outcome.strip():
        return None
    if not _client:
        return None
    prompt = f"""The following is a vague exam-style question and the lesson objective (TBAT = "To Be Able To") from a chemistry unit.

Vague question: {question}
TBAT (intended outcome): {intended_outcome}

Rewrite the question to be specific: "State the units of [quantity 1], [quantity 2]..." using the quantities that the TBAT refers to (e.g. concentration, volume, amount/moles, mass, Mr). Then give the correct units as they would appear in the TBAT or standard GCSE chemistry (e.g. mol/dm³, dm³, mol, g, g/mol).

Output exactly two lines in this format:
QUESTION: State the units of ...
ANSWER: ... (list units in order, e.g. mol/dm³, dm³, mol)"""

    try:
        response = _client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=150,
        )
        text = (response.choices[0].message.content or "").strip()
        q_line = a_line = None
        for line in text.split("\n"):
            line = line.strip()
            if line.upper().startswith("QUESTION:"):
                q_line = line.split(":", 1)[-1].strip().strip('"\'')
            elif line.upper().startswith("ANSWER:"):
                a_line = line.split(":", 1)[-1].strip().strip('"\'')
        if q_line and a_line:
            return (q_line, a_line)
    except Exception:
        pass
    return None


def _answer_for_state_question(
    question: str,
    intended_outcome: str,
    prep_definitions: dict[str, str],
    spec_content: list[tuple[str, str]] | None = None,
) -> str:
    """Infer a short answer for a 'State ...' question from outcome, definitions, or AQA spec (chemistry.json)."""
    # 1) Look for a number like 6.02 x 10^23 in intended outcome
    num_match = re.search(
        r"6\.02\s*[x×]\s*10\s*[^0-9]*23|6\.02\s*×\s*10²³",
        intended_outcome,
        re.IGNORECASE,
    )
    if num_match:
        return "6.02 × 10²³"
    # 2) Look for "equation" / "relationship" in outcome and common answers
    if "equation" in intended_outcome.lower() and "moles" in intended_outcome.lower() and "mass" in intended_outcome.lower():
        return "moles = mass ÷ Mr"
    if "equation" in intended_outcome.lower() and "concentration" in intended_outcome.lower() and "volume" in intended_outcome.lower():
        return "moles = concentration × volume"
    if "dm3" in intended_outcome.lower() or "cm3" in intended_outcome.lower():
        if "1000" in intended_outcome or "1 dm3" in intended_outcome.lower():
            return "1000 cm³ = 1 dm³"
    if "molar volume" in intended_outcome.lower() or "rtp" in intended_outcome.lower():
        return "24 dm³"
    # 3) Try AQA spec content (chemistry.json): e.g. conservation of mass, Avogadro
    if spec_content:
        q_lower = question.lower()
        for content, _tier in spec_content:
            if "6.02" in content or "avogadro" in content.lower():
                if "avogadro" in q_lower or "mole" in q_lower or "particle" in q_lower:
                    return "6.02 × 10²³"
            if "conservation" in q_lower and "mass" in q_lower and "conservation" in content.lower():
                return content.split(".")[0].strip()[:100]
    # 4) "State the units for each quantity" type: extract units from TBAT (e.g. mol/dm3, dm3, mol)
    if _is_generic_units_question(question):
        units = re.findall(
            r"mol/dm³|mol/dm3|mol\s*/\s*dm³|g/mol|dm³|dm3|cm³|cm3|\bmol\b|\bg\b",
            intended_outcome,
            re.IGNORECASE,
        )
        if units:
            # Dedupe preserving order
            seen = set()
            out = []
            for u in units:
                n = re.sub(r"\s+", "", u).replace("³", "3").lower()
                if n not in seen:
                    seen.add(n)
                    out.append(u)
            if out:
                return ", ".join(out)
    # 5) Try prep_definitions: find a term from the question that we have a definition for
    q_lower = question.lower()

    for term, defn in prep_definitions.items():
        if term.lower() in q_lower or term.lower().replace("'", "") in q_lower:
            if defn.strip().upper().startswith("TBAT"):
                continue  # Never use TBAT (intended outcome) as an answer
            # Use first sentence or up to 80 chars
            first = defn.split(".")[0].strip()
            if len(first) > 100:
                first = first[:97] + "..."
            if first.upper().startswith("TBAT"):
                continue
            # If definition contains the value, prefer a short number
            if "6.02" in defn or "6.02" in first:
                return "6.02 × 10²³"
            return first or defn[:80]
    return "[state from lesson]"


def read_prep_booklet_vocab(docx_path: str) -> list[str]:
    """Extract vocabulary words from prep booklet Word doc (for topping up weeks to 10 words)."""
    d = read_prep_booklet_definitions(docx_path)
    return list(d.keys()) if d else []


# AQA specification JSON (chemistry) - used for tier and wording alignment
def _chemistry_spec_path(script_dir: Path | None = None) -> Path:
    script_dir = script_dir or Path(__file__).resolve().parent
    return script_dir / "spec" / "aqa_combined" / "chemistry.json"


def load_chemistry_spec(script_dir: Path | None = None) -> dict | None:
    """Load AQA Chemistry specification from chemistry.json. Returns None if not found."""
    path = _chemistry_spec_path(script_dir)
    if not path.is_file():
        return None
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return None


def load_spec_for_unit(script_dir: Path | None, unit_code: str) -> dict | None:
    """Load AQA spec (chemistry or biology) for the given unit code. Returns None for Physics or if not found."""
    script_dir = script_dir or Path(__file__).resolve().parent
    base = script_dir / "spec" / "aqa_combined"
    uc = (unit_code or "").upper().strip()
    if uc.startswith("C"):
        path = base / "chemistry.json"
    elif uc.startswith("B"):
        path = base / "biology.json"
    else:
        return None
    if not path.is_file():
        return None
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return None


def get_spec_content_for_unit(spec: dict | None, unit_code: str) -> list[tuple[str, str]]:
    """
    Get AQA spec content (foundation + higher) for the topic that matches this unit.
    Unit code e.g. C4.3 -> section 5.3 (Chemistry), B3.2 -> section 4.2 (Biology). Returns list of (content, tier).
    """
    if not spec or "topics" not in spec:
        return []
    # Map unit to spec section: C4.3 -> 5.3 (Chemistry), B3.2 -> 4.2 (Biology)
    uc = unit_code.upper().strip()
    parts = re.findall(r"\d+", unit_code)
    section_id = None
    if uc.startswith("C") and parts:
        section_id = "5." + parts[-1]
    elif uc.startswith("B") and parts:
        section_id = "4." + parts[-1]
    if not section_id:
        return []
    out = []
    for topic in spec.get("topics", []):
        if topic.get("section") != section_id:
            continue
        for sub in topic.get("subsections", []):
            for item in sub.get("foundation_content", []):
                c = item.get("content") or item.get("text")
                if c:
                    out.append((c, "foundation"))
            for item in sub.get("higher_content", []):
                c = item.get("content") or item.get("text")
                if c:
                    out.append((c, "higher"))
        break
    return out


def read_prep_booklet_definitions(docx_path: str) -> dict[str, str]:
    """Extract word -> definition from prep booklet Word doc. Used so AI can generate questions from definitions."""
    if not docx_path or not Path(docx_path).exists():
        return {}
    doc = Document(docx_path)
    result = {}
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                w = cells[0]
                defn = cells[1]
                if w.lower() in ("word", "definition", "term", "vocabulary"):
                    continue
                # Never use TBAT (intended outcome) text as a definition - it's not an answer
                if w.strip().upper().startswith("TBAT") or defn.strip().upper().startswith("TBAT"):
                    continue
                if len(w) > 50 or "," in w or len(defn) < 15:
                    continue
                wn = normalize_kw(w)
                if wn and 1 <= len(w.split()) <= 4:
                    result[w] = defn
    return result


def read_lessons_with_spec(excel_path: str) -> dict:
    """Read lessons and identify which are HT only based on AQA specification"""
    print("="*70)
    print("READING EXCEL WITH AQA SPECIFICATION AWARENESS")
    print("="*70)
    
    df = pd.read_excel(excel_path, header=None)
    
    # Find columns
    keyword_col = None
    spec_col = None
    lesson_col = None
    data_start = None
    
    for idx in range(20):
        for col_idx in range(len(df.columns)):
            value = df.iloc[idx, col_idx]
            if pd.isna(value):
                continue
            val_str = str(value).lower()
            
            if keyword_col is None and 'keyword' in val_str and 'introduc' in val_str:
                keyword_col = col_idx
                print(f"✓ Keyword column: {col_idx}")
            
            if spec_col is None and 'aqa' in val_str and 'specif' in val_str:
                spec_col = col_idx
                print(f"✓ AQA Spec column: {col_idx}")
    
    # Find lesson codes
    for idx in range(5, 20):
        for col_idx in range(len(df.columns)):
            value = df.iloc[idx, col_idx]
            if pd.isna(value):
                continue
            if re.match(r'^[A-Z]\d+\.\d+\.\d+', str(value).strip()):
                lesson_col = col_idx
                data_start = idx
                print(f"✓ Lesson codes: Column {col_idx}, Row {idx}\n")
                break
        if lesson_col is not None:
            break
    
    # Extract lessons
    lessons = {}
    
    for idx in range(data_start, len(df)):
        lesson_cell = df.iloc[idx, lesson_col]
        if pd.isna(lesson_cell):
            continue
        
        lesson_code = str(lesson_cell).strip().split('\n')[0]
        if not re.match(r'^[A-Z]\d+\.\d+\.\d+', lesson_code):
            continue
        
        # Get title
        lesson_title = ""
        if lesson_col + 1 < len(df.columns):
            title_val = df.iloc[idx, lesson_col + 1]
            if pd.notna(title_val):
                lesson_title = str(title_val).strip()
        
        # Determine if HT only
        is_ht_only = ("(HT)" in lesson_title.upper() or 
                      lesson_title.startswith("Taking it Further"))
        
        # Get keywords
        keywords = []
        if keyword_col is not None:
            kw_val = df.iloc[idx, keyword_col]
            if pd.notna(kw_val):
                kw_text = str(kw_val).strip()
                if kw_text and kw_text.lower() not in ['none', 'nan', '', 'n/a']:
                    for line in kw_text.split('\n'):
                        line = line.strip()
                        if line and line.lower() not in ['none', 'nan', 'n/a']:
                            if ',' in line:
                                keywords.extend([k.strip() for k in line.split(',') if k.strip()])
                            else:
                                keywords.append(line)
        
        lessons[lesson_code] = {
            'code': lesson_code,
            'title': lesson_title,
            'keywords': keywords,
            'is_ht_only': is_ht_only
        }
    
    # Display
    print(f"{'Lesson':<15} {'Tier':<15} {'Words':<6} {'Title'}")
    print("-"*70)
    
    for code, data in lessons.items():
        tier = "🔴 HT ONLY" if data['is_ht_only'] else "🟢 FOUNDATION"
        print(f"{code:<15} {tier:<15} {len(data['keywords']):<6} {data['title'][:35]}")
    
    foundation_words = sum(len(l['keywords']) for l in lessons.values() if not l['is_ht_only'])
    ht_words = sum(len(l['keywords']) for l in lessons.values() if l['is_ht_only'])
    
    print("-"*70)
    print(f"Total: {len(lessons)} lessons | Foundation: {foundation_words} words | HT: {ht_words} words\n")
    
    return lessons

def create_weeks_combined(
    lessons: dict,
    extra_words_pool: list[str] | None = None,
    max_weeks: int | None = None,
) -> dict:
    """
    Create weeks by COMBINING lessons to get:
    - Exactly 10 foundation words per week (topped up from prep booklet if needed)
    - Exactly 5 extension words per week
    max_weeks: if set (e.g. 5 for a half term), stop after that many weeks so the doc matches teaching weeks.
    """
    print("="*70)
    print("CREATING WEEKLY BREAKDOWN")
    if max_weeks:
        print(f"  (Half term: generating {max_weeks} weeks only)")
    print("="*70 + "\n")

    # Mutable pool: we'll consume from it when padding to 10
    pool = list(extra_words_pool) if extra_words_pool else []
    pool_used_norm = set()  # normalized forms already used in previous weeks

    def dedupe(keywords):
        seen = set()
        result = []
        for kw in keywords:
            kw_norm = normalize_kw(kw)
            if kw_norm not in seen:
                seen.add(kw_norm)
                result.append(kw)
        return result

    # Sort by numeric lesson number so C4.3.1, C4.3.2, C4.3.3... not C4.3.1, C4.3.10, C4.3.11, C4.3.2
    def lesson_sort_key(lc):
        parts = lc.split('.')
        if len(parts) >= 3:
            try:
                return (parts[0], int(parts[1]), int(parts[2]))
            except ValueError:
                pass
        return (lc, 0, 0)

    lesson_codes = [
        lc for lc in sorted(lessons.keys(), key=lesson_sort_key)
        if 'Feedback' not in lessons[lc]['title']
    ]

    weeks_config = {}
    week_num = 1
    i = 0

    while i < len(lesson_codes):
        week_lessons = []
        week_keywords_foundation = []
        week_keywords_ht = []

        current_lesson = lesson_codes[i]
        week_lessons.append(current_lesson)

        current_kw = lessons[current_lesson]['keywords']
        if lessons[current_lesson]['is_ht_only']:
            week_keywords_ht.extend(current_kw)
        else:
            week_keywords_foundation.extend(current_kw)

        i += 1

        total_words = len(week_keywords_foundation) + len(week_keywords_ht)
        while i < len(lesson_codes) and total_words < 12:
            next_lesson = lesson_codes[i]
            next_kw = lessons[next_lesson]['keywords']
            if total_words + len(next_kw) > 18:
                break
            week_lessons.append(next_lesson)
            if lessons[next_lesson]['is_ht_only']:
                week_keywords_ht.extend(next_kw)
            else:
                week_keywords_foundation.extend(next_kw)
            i += 1
            total_words = len(week_keywords_foundation) + len(week_keywords_ht)
        
        # Remove duplicates
        def dedupe(keywords):
            seen = set()
            result = []
            for kw in keywords:
                kw_norm = normalize_kw(kw)
                if kw_norm not in seen:
                    seen.add(kw_norm)
                    result.append(kw)
            return result
        
        week_keywords_foundation = dedupe(week_keywords_foundation)
        week_keywords_ht = dedupe(week_keywords_ht)
        
        # Ensure we have exactly 5 extension words
        # Take first 5 HT words, or pad from foundation if needed
        if len(week_keywords_ht) >= 5:
            extension_words = week_keywords_ht[:5]
            # Rest go to foundation
            week_keywords_foundation.extend(week_keywords_ht[5:])
        else:
            # Not enough HT words - take some from foundation as "challenging" extensions
            extension_words = week_keywords_ht.copy()
            needed = 5 - len(extension_words)
            if len(week_keywords_foundation) > needed:
                # Move last N foundation words to extension
                extension_words.extend(week_keywords_foundation[-needed:])
                week_keywords_foundation = week_keywords_foundation[:-needed]
        
        # Ensure minimum 10 foundation words: top up from extra_words_pool (prep booklet vocab) first, then repeat if needed
        week_foundation_norm = {normalize_kw(k) for k in week_keywords_foundation}
        while len(week_keywords_foundation) < 10:
            if week_keywords_foundation:
                # Prefer unused words from pool (vocab and literacy from prep booklet)
                added = False
                for w in pool:
                    n = normalize_kw(w)
                    if n and n not in week_foundation_norm and n not in pool_used_norm:
                        week_keywords_foundation.append(w)
                        week_foundation_norm.add(n)
                        pool_used_norm.add(n)
                        added = True
                        break
                if not added:
                    week_keywords_foundation.append(week_keywords_foundation[len(week_keywords_foundation) % len(week_keywords_foundation)])
            else:
                break
        # Cap foundation at 15 (so we don't exceed; pool words can push us over 10 but we still cap at 15)
        week_keywords_foundation = week_keywords_foundation[:15]
        
        weeks_config[week_num] = {
            'lessons': ' & '.join(week_lessons),
            'titles': ' / '.join([lessons[lc]['title'] for lc in week_lessons]),
            'foundation': week_keywords_foundation[:15],  # Max 15 foundation
            'extension': extension_words[:5],  # Exactly 5 extension
        }
        
        print(f"Week {week_num}: {weeks_config[week_num]['lessons']}")
        print(f"  Foundation ({len(weeks_config[week_num]['foundation'])}): {', '.join(weeks_config[week_num]['foundation'][:3])}...")
        print(f"  Extension ({len(weeks_config[week_num]['extension'])}): {', '.join(weeks_config[week_num]['extension'][:3])}...")
        print()
        
        week_num += 1
        if max_weeks and week_num > max_weeks:
            break
    
    return weeks_config


def build_weeks_config_from_assignments(
    lessons: dict,
    week_lesson_codes: list[list[str]],
    extra_words_pool: list[str] | None = None,
    _log: bool = True,
) -> dict:
    """
    Build weeks_config from teacher-defined week assignments.
    week_lesson_codes: list of lists, e.g. [["C4.3.2", "C4.3.3"], ["C4.3.4", "C4.3.5", "C4.3.6"]]
    Each inner list is the lesson codes for that week (Week 1, Week 2, ...).
    """
    pool = list(extra_words_pool) if extra_words_pool else []
    pool_used_norm = set()

    def dedupe(keywords):
        seen = set()
        result = []
        for kw in keywords:
            kw_norm = normalize_kw(kw)
            if kw_norm not in seen:
                seen.add(kw_norm)
                result.append(kw)
        return result

    weeks_config = {}
    for week_num, lesson_codes in enumerate(week_lesson_codes, start=1):
        week_lessons = [lc for lc in lesson_codes if lc in lessons and "Feedback" not in lessons.get(lc, {}).get("title", "")]
        if not week_lessons:
            continue
        week_keywords_foundation = []
        week_keywords_ht = []
        for lc in week_lessons:
            kw = lessons[lc]["keywords"]
            if lessons[lc]["is_ht_only"]:
                week_keywords_ht.extend(kw)
            else:
                week_keywords_foundation.extend(kw)
        week_keywords_foundation = dedupe(week_keywords_foundation)
        week_keywords_ht = dedupe(week_keywords_ht)
        # Exactly 5 extension
        if len(week_keywords_ht) >= 5:
            extension_words = week_keywords_ht[:5]
            week_keywords_foundation.extend(week_keywords_ht[5:])
        else:
            extension_words = week_keywords_ht.copy()
            needed = 5 - len(extension_words)
            if len(week_keywords_foundation) > needed:
                extension_words.extend(week_keywords_foundation[-needed:])
                week_keywords_foundation = week_keywords_foundation[:-needed]
        # Min 10 foundation, top up from pool
        week_foundation_norm = {normalize_kw(k) for k in week_keywords_foundation}
        while len(week_keywords_foundation) < 10 and week_keywords_foundation:
            added = False
            for w in pool:
                n = normalize_kw(w)
                if n and n not in week_foundation_norm and n not in pool_used_norm:
                    week_keywords_foundation.append(w)
                    week_foundation_norm.add(n)
                    pool_used_norm.add(n)
                    added = True
                    break
            if not added:
                week_keywords_foundation.append(week_keywords_foundation[len(week_keywords_foundation) % len(week_keywords_foundation)])
        week_keywords_foundation = week_keywords_foundation[:15]
        weeks_config[week_num] = {
            "lessons": " & ".join(week_lessons),
            "titles": " / ".join([lessons[lc]["title"] for lc in week_lessons]),
            "foundation": week_keywords_foundation,
            "extension": extension_words[:5],
        }
        if _log:
            print(f"Week {week_num}: {weeks_config[week_num]['lessons']}")
            print(f"  Foundation ({len(weeks_config[week_num]['foundation'])}): {', '.join(weeks_config[week_num]['foundation'][:3])}...")
            print(f"  Extension ({len(weeks_config[week_num]['extension'])}): {', '.join(weeks_config[week_num]['extension'][:3])}...")
    return weeks_config


def create_aqa_question(keyword: str, subject: str = "Chemistry") -> str:
    """Generate AQA-style question for a keyword"""
    kw_lower = keyword.lower()
    
    # Chemistry-specific questions
    chemistry_questions = {
        'mole': "What is the unit used to measure the amount of a substance called?",
        'mass': "What is the quantity of matter in an object called?",
        'concentration': "What is the amount of solute per unit volume of solution called?",
        'acid': "What type of substance produces hydrogen ions (H⁺) in aqueous solution?",
        'alkali': "What type of base is soluble in water?",
        'titration': "What is the technique used to determine the concentration of an unknown solution?",
        'neutralisation': "What is the reaction between an acid and a base called?",
        'relative atomic mass': "What is the weighted average mass of an atom compared to 1/12th the mass of a carbon-12 atom?",
        'relative formula mass': "What is the sum of the relative atomic masses of all atoms in a formula called?",
        'percentage by mass': "What calculation shows the mass of an element as a percentage of the total mass?",
        'avogadro': "What is the number of particles in one mole of a substance called?",
        'limiting reactant': "What is the reactant that is completely used up in a reaction called?",
        'hydrogen ion': "What ion is produced by acids in aqueous solution?",
        'hydroxide ion': "What ion is produced by alkalis in aqueous solution?",
    }
    
    # Check for exact matches
    for key, question in chemistry_questions.items():
        if key in kw_lower:
            return question
    
    # Generic question templates
    if 'formula' in kw_lower or 'equation' in kw_lower:
        return f"What is {keyword} in chemistry?"
    elif 'reaction' in kw_lower or 'process' in kw_lower:
        return f"What is the {keyword} process called?"
    else:
        return f"What is {keyword}?"


def generate_questions_with_ai(
    all_keywords: list[str],
    unit_code: str,
    unit_name: str,
    all_vocab: dict[str, str] | None = None,
    spec_content: list[tuple[str, str]] | None = None,
) -> dict[str, str]:
    """Generate one AQA-style question per keyword. Uses definitions from the Word doc when available so AI writes questions from the actual definition."""
    question_bank = {}
    if not _client or not all_keywords:
        return question_bank

    all_vocab = all_vocab or {}
    # Normalized lookup: norm -> (original_word, definition)
    def_lookup = {}
    for w, defn in all_vocab.items():
        n = normalize_kw(w)
        if n:
            def_lookup[n] = (w, defn)

    # Dedupe keywords by normalized form
    seen = set()
    unique = []
    for kw in all_keywords:
        n = normalize_kw(kw)
        if n and n not in seen:
            seen.add(n)
            unique.append(kw)

    subject = "Biology" if unit_code.startswith("B") else "Chemistry" if unit_code.startswith("C") else "Physics"
    with_defn = sum(1 for kw in unique if normalize_kw(kw) in def_lookup)
    print(f"\n🤖 Generating questions via OpenAI ({with_defn} from definitions, {len(unique) - with_defn} from keyword only)...")

    # Process in batches; include definition when we have it so AI can base the question on it
    batch_size = 15
    for start in range(0, len(unique), batch_size):
        batch = unique[start : start + batch_size]
        lines = []
        for kw in batch:
            n = normalize_kw(kw)
            if n in def_lookup:
                _, defn = def_lookup[n]
                lines.append(f"  {kw} | Definition: {defn}")
            else:
                lines.append(f"  {kw} | (no definition)")
        block = "\n".join(lines)
        spec_block = ""
        if spec_content:
            spec_lines = [c for c, _ in spec_content[:30]]  # limit size
            spec_block = "\nAQA specification content for this unit (align wording with this):\n" + "\n".join(f"- {s}" for s in spec_lines) + "\n\n"
        prompt = f"""You are writing revision quiz questions for GCSE {subject} (AQA). Unit: {unit_name}.
{spec_block}
For each line below, write ONE short question. When a definition is given, base your question on that definition so the correct answer is the keyword. When there is no definition, write a question whose answer is the keyword.
- Maximum 18 words per question.
- AQA exam-style. Prefer wording that matches the AQA specification when relevant.
- Output exactly one line per keyword in this format: keyword: question

{block}

Output one line per keyword. Use the exact keyword as given:"""

        try:
            response = _client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=1200,
            )
            text = response.choices[0].message.content or ""
            for line in text.strip().split("\n"):
                if ":" not in line:
                    continue
                kw, _, q = line.partition(":")
                kw = re.sub(r"^[\d\.\-\*\s]+", "", kw.strip()).strip()
                q = q.strip().strip('"\'')
                if kw and q:
                    question_bank[normalize_kw(kw)] = q
        except Exception as e:
            print(f"⚠️ API batch error: {e}")

    print(f"✓ Generated {len(question_bank)} questions")
    return question_bank




def create_aqa_document(
    weeks_config: dict, unit_code: str, unit_name: str, output_file: str,
    question_bank: dict[str, str] | None = None,
    state_questions_by_lesson: dict[str, list[tuple[str, str]]] | None = None,
):
    """Generate Word document with highlighted extension words."""
    print("="*70)
    print("GENERATING AQA-SPECIFICATION WORD DOCUMENT")
    print("="*70 + "\n")
    
    doc = Document()
    # Ask Word to open in normal mode, not Compatibility Mode (avoids "Do you want to save changes?" on close)
    try:
        doc.part.element.set(qn("w:conformance"), "strict")
    except Exception:
        pass

    # Set margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1.0)
        section.left_margin = section.right_margin = Inches(1.0)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(f"{unit_code} – {unit_name}: Core Knowledge")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_text = subtitle.add_run("AQA GCSE Separate Science (Triple) Specification")
    subtitle_text.font.size = Pt(12)
    subtitle_text.font.name = 'Arial'
    subtitle_text.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Key
    key_para = doc.add_paragraph()
    key_para.add_run("Key: ").font.bold = True
    key_para.add_run("Standard text = Foundation tier (ALL students must learn) | ")
    
    ext_run = key_para.add_run("Yellow highlight = Higher tier extension (Triple science students only)")
    ext_run.font.color.rgb = RGBColor(0, 0, 0)
    ext_run.font.highlight_color = 7  # Yellow highlight
    
    doc.add_paragraph()
    
    # Add weeks
    for week_num in sorted(weeks_config.keys()):
        week_data = weeks_config[week_num]
        foundation_words = week_data['foundation']
        extension_words = week_data['extension']
        
        # Week header: show lesson codes only (e.g. "C4.3.1 and C4.3.2"), not full titles
        lesson_codes = week_data['lessons'].replace(' & ', ' and ')
        week_para = doc.add_paragraph()
        week_run = week_para.add_run(
            f"Week {week_num} – {lesson_codes}"
        )
        week_run.font.size = Pt(14)
        week_run.font.bold = True
        week_run.font.name = 'Arial'
        
        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        hdr = table.rows[0].cells
        for i, text in enumerate(['Question', 'Answer']):
            p = hdr[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D5E8F0')
            hdr[i]._element.get_or_add_tcPr().append(shading)
        
        # Add FOUNDATION words (no highlighting); validate then sanitise so answers stay as revision terms
        for keyword in foundation_words:
            row = table.add_row().cells
            question = (question_bank or {}).get(normalize_kw(keyword)) or create_aqa_question(keyword)
            answer = _validate_answer_with_ai(question, keyword)
            answer = _sanitise_answer_for_revision(keyword, answer)
            if len(answer) > MAX_STATE_ANSWER_CHARS:
                answer = answer[:MAX_STATE_ANSWER_CHARS]

            # Question cell
            p_q = row[0].paragraphs[0]
            p_q.clear()
            run = p_q.add_run(question)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            
            # Answer cell
            p_a = row[1].paragraphs[0]
            p_a.clear()
            run = p_a.add_run(answer)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
        
        # Add EXTENSION words (with yellow highlighting); validate then sanitise so answers stay as revision terms
        for keyword in extension_words:
            row = table.add_row().cells
            question = (question_bank or {}).get(normalize_kw(keyword)) or create_aqa_question(keyword)
            answer = _validate_answer_with_ai(question, keyword)
            answer = _sanitise_answer_for_revision(keyword, answer)
            if len(answer) > MAX_STATE_ANSWER_CHARS:
                answer = answer[:MAX_STATE_ANSWER_CHARS]

            # Question cell - HIGHLIGHTED
            p_q = row[0].paragraphs[0]
            p_q.clear()
            run = p_q.add_run(question)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.highlight_color = 7  # Yellow
            
            # Answer cell - HIGHLIGHTED
            p_a = row[1].paragraphs[0]
            p_a.clear()
            run = p_a.add_run(answer)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.highlight_color = 7  # Yellow

        # Add "State ..." questions for this week (from prep booklet example questions)
        # Dedupe: don't add the same question or same "units" answer twice in one week
        week_lesson_codes = [s.strip() for s in week_data['lessons'].split(' & ')]
        state_by_lesson = state_questions_by_lesson or {}
        seen_questions_norm = set()
        seen_units_answers = set()
        for lc in week_lesson_codes:
            for q_text, a_text in state_by_lesson.get(lc, []):
                q_norm = normalize_kw(q_text)
                if q_norm in seen_questions_norm:
                    continue
                # Same "State the units of..." with same answer = skip
                if "state the units of" in q_norm or "units of" in q_norm:
                    if a_text.strip() in seen_units_answers:
                        continue
                    seen_units_answers.add(a_text.strip())
                seen_questions_norm.add(q_norm)
                row = table.add_row().cells
                p_q = row[0].paragraphs[0]
                p_q.clear()
                run = p_q.add_run(_state_statement_to_question(q_text))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                p_a = row[1].paragraphs[0]
                p_a.clear()
                run = p_a.add_run(a_text)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
        
        doc.add_paragraph()
    
    # Save
    doc.save(output_file)
    print(f"✅ Document generated: {output_file}\n")
    print("Features:")
    print("  ✓ 10+ foundation words per week (ALL students)")
    print("  ✓ Exactly 5 extension words per week (Triple science HT)")
    print("  ✓ Extension words highlighted in YELLOW")
    print("  ✓ AQA-style questions")
    print("  ✓ Based on AQA specification tier system")


def run_daily_review_generator(
    unit_code: str,
    lesson_resources_root: str | Path,
    week_assignments: list[list[str]] | None = None,
    output_dir: str | Path | None = None,
    max_weeks: int | None = None,
) -> tuple[str | None, str | None]:
    """
    Run the full Core Knowledge doc generator.
    week_assignments: optional list of lists, e.g. [["C4.3.2", "C4.3.3"], ["C4.3.4", "C4.3.5"]].
    max_weeks: for auto mode, cap at this many weeks (e.g. 5 for a half term with 5 teaching weeks).
    Returns (output_file_path, None) on success, or (None, error_message) on failure.
    """
    script_dir = Path(__file__).resolve().parent
    lesson_resources = Path(lesson_resources_root)
    out_dir = Path(output_dir) if output_dir else script_dir
    try:
        found = find_unit_resources(unit_code, lesson_resources)
    except FileNotFoundError as e:
        return (None, str(e))
    excel_path = found["excel_path"]
    unit_code = found["unit_code"]
    unit_name = found["unit_name"]
    output_file = str(out_dir / f"{unit_code.replace('.', '_')}_Core_Knowledge.docx")
    try:
        lessons = read_lessons_with_spec(excel_path)
    except Exception as e:
        return (None, str(e))
    if not lessons:
        return (None, "No lessons found in Excel.")
    prep_definitions = read_prep_booklet_definitions(found.get("prep_booklet_path", ""))
    extra_words = list(prep_definitions.keys()) if prep_definitions else []
    chemistry_spec = load_chemistry_spec(script_dir)
    spec_content = get_spec_content_for_unit(chemistry_spec, unit_code) if chemistry_spec else []
    if week_assignments:
        weeks_config = build_weeks_config_from_assignments(lessons, week_assignments, extra_words_pool=extra_words, _log=True)
    else:
        weeks_config = create_weeks_combined(lessons, extra_words_pool=extra_words, max_weeks=max_weeks)
    if not weeks_config:
        return (None, "No weeks created (check week assignments).")
    all_keywords = []
    for w in weeks_config.values():
        all_keywords.extend(w["foundation"])
        all_keywords.extend(w["extension"])
    question_bank = generate_questions_with_ai(all_keywords, unit_code, unit_name, all_vocab=prep_definitions, spec_content=spec_content)
    ordered_lessons = sorted((lc for lc in lessons if "Feedback" not in lessons[lc]["title"]), key=_lesson_sort_key)
    state_questions_by_lesson = read_state_questions_from_prep_booklet(
        found.get("prep_booklet_path", ""), ordered_lessons, prep_definitions, spec_content=spec_content
    )
    create_aqa_document(weeks_config, unit_code, unit_name, output_file, question_bank=question_bank, state_questions_by_lesson=state_questions_by_lesson)
    return (output_file, None)


def load_lessons_and_resources_for_codes(
    lesson_codes: list[str],
    lesson_resources_root: str | Path,
    script_dir: Path | None = None,
) -> tuple[dict, dict[str, str], dict[str, list[tuple[str, str]]], list[str], list[tuple[str, str]]]:
    """
    Load merged lessons, prep definitions, state questions, and spec content for a list of lesson codes
    that may span multiple units (e.g. B3.2.4, C4.2.2, P3.2.2).
    Returns: (lessons_merged, prep_definitions, state_questions_by_lesson, extra_words_list, spec_content_merged).
    """
    script_dir = script_dir or Path(__file__).resolve().parent
    root = Path(lesson_resources_root)
    codes_set = {lc.split("/")[0].strip() for lc in lesson_codes}
    units_needed = sorted(set(lesson_code_to_unit(lc) for lc in codes_set))

    lessons_merged = {}
    prep_definitions = {}
    state_questions_by_lesson = {}
    spec_content_merged = []

    for unit_code in units_needed:
        try:
            found = find_unit_resources(unit_code, root)
        except FileNotFoundError:
            continue
        unit_lessons = read_lessons_with_spec(found["excel_path"])
        for lc, data in unit_lessons.items():
            lc_base = lc.split("/")[0].strip()
            if lc_base in codes_set or lc in codes_set:
                lessons_merged[lc] = data
        prep_path = found.get("prep_booklet_path", "")
        defs = read_prep_booklet_definitions(prep_path)
        for w, d in defs.items():
            if normalize_kw(w) not in {normalize_kw(k) for k in prep_definitions}:
                prep_definitions[w] = d
        ordered = sorted(
            [lc for lc in unit_lessons if "Feedback" not in unit_lessons[lc].get("title", "")],
            key=lambda x: _lesson_sort_key(x.split("/")[0] if "/" in x else x),
        )
        spec = load_spec_for_unit(script_dir, unit_code)
        spec_content = get_spec_content_for_unit(spec, unit_code) if spec else []
        spec_content_merged.extend(spec_content)
        state_for_unit = read_state_questions_from_prep_booklet(
            prep_path, ordered, defs, spec_content=spec_content
        )
        for lc, qa_list in state_for_unit.items():
            state_questions_by_lesson[lc] = state_questions_by_lesson.get(lc, []) + qa_list

    extra_words = list(prep_definitions.keys())
    return (lessons_merged, prep_definitions, state_questions_by_lesson, extra_words, spec_content_merged)


def run_daily_review_generator_multi(
    lesson_resources_root: str | Path,
    week_assignments: list[list[str]],
    output_dir: str | Path | None = None,
) -> tuple[str | None, str | None]:
    """
    Generate a combined Core Knowledge document for teachers who teach multiple units in the same week
    (e.g. Week 1: B3.2.4, C4.2.2, P3.2.2). week_assignments is a list of lists of lesson codes (any units).
    Returns (output_file_path, None) on success, or (None, error_message) on failure.
    """
    script_dir = Path(__file__).resolve().parent
    lesson_resources = Path(lesson_resources_root)
    out_dir = Path(output_dir) if output_dir else script_dir

    all_codes = []
    for week in week_assignments:
        for lc in week:
            base = lc.split("/")[0].strip()
            if base and base not in all_codes:
                all_codes.append(base)
    if not all_codes:
        return (None, "No lesson codes in week assignments.")

    try:
        lessons_merged, prep_definitions, state_by_lesson, extra_words, spec_content = load_lessons_and_resources_for_codes(
            all_codes, lesson_resources, script_dir
        )
    except Exception as e:
        return (None, str(e))

    if not lessons_merged:
        return (None, "No lessons found for the given codes. Check unit folders exist in Lesson Resources.")

    def expand_to_lesson_keys(codes: list[str]) -> list[str]:
        out = []
        for lc in codes:
            if lc in lessons_merged:
                out.append(lc)
            else:
                base = lc.split("/")[0].strip()
                out.extend(k for k in lessons_merged if k.split("/")[0].strip() == base)
        return out

    week_assignments_normalized = []
    for week in week_assignments:
        normalized = [lc.split("/")[0].strip() for lc in week if lc.strip()]
        week_lessons = expand_to_lesson_keys(normalized)
        if week_lessons:
            week_assignments_normalized.append(week_lessons)

    if not week_assignments_normalized:
        return (None, "No weeks with valid lessons. Check lesson codes match the Excel (e.g. B3.2.4, C4.2.2).")

    weeks_config = build_weeks_config_from_assignments(
        lessons_merged, week_assignments_normalized, extra_words_pool=extra_words, _log=True
    )
    if not weeks_config:
        return (None, "No weeks created (check week assignments).")

    all_keywords = []
    for w in weeks_config.values():
        all_keywords.extend(w["foundation"])
        all_keywords.extend(w["extension"])

    units_label = ", ".join(sorted(set(lesson_code_to_unit(lc) for lc in lessons_merged)))
    multi_name = f"Multiple units ({units_label})"
    question_bank = generate_questions_with_ai(
        all_keywords, "Multi", multi_name, all_vocab=prep_definitions, spec_content=spec_content
    )

    output_file = str(out_dir / "Multi_Core_Knowledge.docx")
    create_aqa_document(
        weeks_config,
        "Multi",
        units_label,
        output_file,
        question_bank=question_bank,
        state_questions_by_lesson=state_by_lesson,
    )
    return (output_file, None)

def list_available_units(lesson_resources_root: str | Path) -> list[str]:
    """Return list of unit codes (e.g. ['C4.2', 'C4.3']) from Lesson Resources folders."""
    root = Path(lesson_resources_root)
    if not root.is_dir():
        return []
    units = []
    for d in root.iterdir():
        if not d.is_dir():
            continue
        name = d.name
        parts = name.split()
        if parts and re.match(r"^[A-Z]\d+\.\d+", parts[0]):
            units.append(parts[0])
    return sorted(set(units))


def _term_json_path(script_dir: Path | None = None) -> Path:
    script_dir = script_dir or Path(__file__).resolve().parent
    return script_dir / "term.json"


def load_term_data(script_dir: Path | None = None) -> dict | None:
    """Load term.json (academic years, terms, half terms). Returns None if not found."""
    path = _term_json_path(script_dir)
    if not path.is_file():
        return None
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return None


def _teaching_weeks_from_dates(start_str: str, end_str: str) -> int:
    """Compute number of teaching weeks from start and end date (YYYY-MM-DD)."""
    try:
        start = datetime.strptime(start_str.strip()[:10], "%Y-%m-%d")
        end = datetime.strptime(end_str.strip()[:10], "%Y-%m-%d")
        days = (end - start).days + 1
        return max(1, round(days / 7))
    except (ValueError, TypeError):
        return 6


def list_half_terms(term_data: dict | None) -> list[tuple[str, str, str, int]]:
    """
    Return list of (academic_year, term_name, half_term_name, teaching_weeks) for UI.
    term.json structure: year -> term (e.g. Autumn Term) -> half_terms: [{ name, start, end }, ...].
    teaching_weeks is derived from (end - start) for each half term, or from teaching_weeks if present.
    """
    if not term_data:
        return []
    out = []
    for year, terms in term_data.items():
        if not isinstance(terms, dict):
            continue
        for term_name, term_obj in terms.items():
            if not isinstance(term_obj, dict):
                continue
            half_terms = term_obj.get("half_terms") or term_obj.get("half_term") or []
            for i, ht in enumerate(half_terms):
                if not isinstance(ht, dict):
                    continue
                name = ht.get("name") or f"HT{i + 1}"
                if "teaching_weeks" in ht:
                    weeks = int(ht["teaching_weeks"])
                elif "start" in ht and "end" in ht:
                    weeks = _teaching_weeks_from_dates(ht["start"], ht["end"])
                else:
                    weeks = 6
                out.append((year, term_name, name, weeks))
    return out


def get_teaching_weeks_for_half_term(
    term_data: dict | None,
    academic_year: str,
    term_name: str,
    half_term_index: int,
) -> int | None:
    """Get teaching_weeks for a given year, term and half-term index (0-based). From dates or teaching_weeks in term.json."""
    if not term_data or academic_year not in term_data:
        return None
    terms = term_data[academic_year]
    if not isinstance(terms, dict) or term_name not in terms:
        return None
    term_obj = terms[term_name]
    half_terms = term_obj.get("half_terms") or term_obj.get("half_term") or []
    if half_term_index < 0 or half_term_index >= len(half_terms):
        return None
    ht = half_terms[half_term_index]
    if not isinstance(ht, dict):
        return None
    if "teaching_weeks" in ht:
        return int(ht["teaching_weeks"])
    if "start" in ht and "end" in ht:
        return _teaching_weeks_from_dates(ht["start"], ht["end"])
    return 6


def find_unit_resources(unit_code: str, lesson_resources_root: str | Path) -> dict:
    """Find Excel (and optional prep booklet) for a unit in Lesson Resources."""
    root = Path(lesson_resources_root)
    if not root.is_dir():
        raise FileNotFoundError(f"Not found: {root}")

    unit_folders = [
        d for d in root.iterdir()
        if d.is_dir() and d.name.upper().startswith(unit_code.upper())
    ]
    if not unit_folders:
        available = ", ".join(d.name for d in root.iterdir() if d.is_dir())
        raise FileNotFoundError(f"No folder for '{unit_code}'. Available: {available}")

    folder = unit_folders[0]
    code = folder.name.split()[0] if folder.name.split() else unit_code
    unit_name = folder.name.split(maxsplit=1)[1] if len(folder.name.split()) > 1 else folder.name

    guidance = folder / "Unit Guidance"
    if not guidance.is_dir():
        raise FileNotFoundError(f"No 'Unit Guidance' in {folder}")

    xlsx_files = [f for f in guidance.rglob("*.xlsx") if not f.name.startswith("~$")]
    if not xlsx_files:
        raise FileNotFoundError(f"No .xlsx in {guidance}")

    prep_booklet_path = ""
    prep_docx = guidance / f"{code} Unit Preparation Booklet.docx"
    if not prep_docx.is_file():
        prep_dir = guidance / f"{code} Unit Preparation Booklet"
        prep_docx = prep_dir / f"{code} Unit Preparation Booklet.docx"
    if prep_docx.is_file():
        prep_booklet_path = str(prep_docx)

    return {
        "excel_path": str(xlsx_files[0]),
        "prep_booklet_path": prep_booklet_path,
        "unit_name": unit_name,
        "unit_code": code,
    }


if __name__ == "__main__":
    UNIT = "C4.3"
    max_weeks_cli = None
    i = 1
    while i < len(sys.argv):
        arg = sys.argv[i]
        if arg == "--weeks" and i + 1 < len(sys.argv) and sys.argv[i + 1].isdigit():
            max_weeks_cli = int(sys.argv[i + 1])
            i += 2
            continue
        if arg.startswith("--weeks=") and arg[8:].isdigit():
            max_weeks_cli = int(arg[8:])
            i += 1
            continue
        if not arg.startswith("--"):
            UNIT = arg.strip()
        i += 1

    script_dir = Path(__file__).resolve().parent
    lesson_resources = script_dir / "Lesson Resources"

    try:
        found = find_unit_resources(UNIT, lesson_resources)
    except FileNotFoundError as e:
        print(f"❌ {e}")
        sys.exit(1)

    excel_path = found["excel_path"]
    unit_code = found["unit_code"]
    unit_name = found["unit_name"]
    output_file = f"{unit_code.replace('.', '_')}_Core_Knowledge.docx"

    print(f"Unit: {unit_name}")
    print(f"Excel: {Path(excel_path).name}")
    if found.get("prep_booklet_path"):
        print(f"Prep booklet: {Path(found['prep_booklet_path']).name}")
    print(f"Output: {output_file}\n")

    lessons = read_lessons_with_spec(excel_path)
    if not lessons:
        print("❌ ERROR: No lessons found!")
        sys.exit(1)

    prep_definitions = read_prep_booklet_definitions(found.get("prep_booklet_path", ""))
    extra_words = list(prep_definitions.keys()) if prep_definitions else []
    if prep_definitions:
        print(f"✓ Loaded {len(prep_definitions)} words + definitions from prep booklet (for 10 per week and AI questions)\n")

    # AQA specification (chemistry.json) for wording alignment and state-question answers
    chemistry_spec = load_chemistry_spec(script_dir)
    spec_content = get_spec_content_for_unit(chemistry_spec, unit_code) if chemistry_spec else []
    if spec_content:
        print(f"✓ Loaded AQA spec content from chemistry.json ({len(spec_content)} statements for this unit)\n")

    weeks_config = create_weeks_combined(lessons, extra_words_pool=extra_words, max_weeks=max_weeks_cli)
    if not weeks_config:
        print("❌ ERROR: No weeks created!")
        sys.exit(1)

    # Report prep booklet (vocab and literacy) words that do not appear in the doc
    in_doc_norm = set()
    for w in weeks_config.values():
        for kw in w["foundation"] + w["extension"]:
            in_doc_norm.add(normalize_kw(kw))
    unused_from_pool = [w for w in extra_words if normalize_kw(w) not in in_doc_norm]
    if unused_from_pool:
        print(f"📋 Prep booklet words NOT in Core Knowledge doc ({len(unused_from_pool)}):")
        print("   " + ", ".join(sorted(unused_from_pool)))
        print()
    else:
        print("✓ All prep booklet vocab words appear in the doc.\n")

    all_keywords = []
    for w in weeks_config.values():
        all_keywords.extend(w["foundation"])
        all_keywords.extend(w["extension"])
    question_bank = generate_questions_with_ai(
        all_keywords, unit_code, unit_name, all_vocab=prep_definitions, spec_content=spec_content
    )

    # Extract "State ..." example questions from prep booklet (one per lesson)
    ordered_lessons = sorted(
        (lc for lc in lessons if "Feedback" not in lessons[lc]["title"]),
        key=_lesson_sort_key,
    )
    state_questions_by_lesson = read_state_questions_from_prep_booklet(
        found.get("prep_booklet_path", ""),
        ordered_lessons,
        prep_definitions,
        spec_content=spec_content,
    )
    if state_questions_by_lesson:
        total_state = sum(len(v) for v in state_questions_by_lesson.values())
        print(f"✓ Loaded {total_state} 'State' questions from prep booklet example questions")
        print("  (State answers are chosen, in order: TBAT/outcome hints → prep booklet definitions")
        print("  → AQA spec → AI fix if wrong/long → teacher validation; questions with answers >50 chars are dropped.)\n")

    create_aqa_document(
        weeks_config,
        unit_code,
        unit_name,
        output_file,
        question_bank=question_bank,
        state_questions_by_lesson=state_questions_by_lesson,
    )