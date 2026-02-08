import openai
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
import json
import re
import time

openai.api_key = os.getenv('OPENAI_API_KEY')

class LessonContentExtractor:
    """Handles reading and extracting lesson data from Word/Excel files"""
    
    @staticmethod
    def read_word_doc(file_path):
        try:
            doc = Document(file_path)
            return '\n'.join([p.text for p in doc.paragraphs])
        except Exception as e:
            print(f"Error reading document: {e}")
            return ""
    
    @staticmethod
    def read_excel_objectives(file_path):
        try:
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active
            objectives = {}
            
            for row in sheet.iter_rows(values_only=True):
                if row and len(row) >= 3 and isinstance(row[0], str) and row[0].startswith('C4.1.'):
                    lesson_code = row[0].strip()
                    lesson_num = lesson_code.split('.')[-1]
                    
                    objectives[lesson_num] = {
                        'lesson_number': lesson_num,
                        'title': str(row[1]) if row[1] else f"Lesson {lesson_num}",
                        'know': [s.strip() for s in str(row[2]).split('\n') if s.strip()] if row[2] else [],
                        'do': [s.strip() for s in str(row[4]).split('\n') if s.strip()] if len(row) > 4 and row[4] else [],
                        'common_errors': str(row[9]) if len(row) > 9 and row[9] else ""
                    }
            
            return objectives
        except Exception as e:
            print(f"Error reading Excel: {e}")
            return {}
    
    @staticmethod
    def extract_lessons(document_text):
        """Extract lesson titles from document"""
        lessons = []
        lines = document_text.split('\n')
        
        for i, line in enumerate(lines):
            match = re.search(r'lesson\s+(\d+)[:\s]*(.*)', line.strip(), re.IGNORECASE)
            if match:
                lesson_num = match.group(1)
                lesson_title = match.group(2).strip() or (lines[i + 1].strip() if i + 1 < len(lines) else "")
                
                lessons.append({
                    'lesson_number': lesson_num,
                    'title': f"Lesson {lesson_num}: {lesson_title}" if lesson_title else f"Lesson {lesson_num}"
                })
        
        return lessons if lessons else LessonContentExtractor.get_default_lessons()
    
    @staticmethod
    def get_default_lessons():
        """Fallback lessons if extraction fails"""
        return [
            {'lesson_number': '1', 'title': 'Covalent Bonding'},
            {'lesson_number': '2', 'title': 'Ionic Bonding'},
            {'lesson_number': '3', 'title': 'Metallic Bonding'}
        ]


class AIContentGenerator:
    """Handles AI content generation via OpenAI"""
    
    @staticmethod
    def generate_lesson_metadata(lesson_title):
        """Generate topic, example, and learning objectives"""
        prompt = f"""Generate simple chemistry lesson content for Year 10 GCSE students (ages 14-15).
Lesson: "{lesson_title}"

IMPORTANT: Use ONLY simple GCSE-level language. NO advanced terms.

Return ONLY valid JSON:
{{
    "topic": "Simple topic name (2-4 words)",
    "example": "Chemical example (e.g., H₂O, NaCl)",
    "tkt_statements": [
        "5 simple knowledge statements using basic language",
        "Avoid university words - use everyday language",
        "Keep each statement to 10-15 words maximum"
    ],
    "tbat_statements": [
        "2 simple skill statements students can do",
        "Use action verbs: draw, explain, describe, identify"
    ],
    "common_errors": "Simple description of what students get wrong and how to fix it (2-3 sentences)"
}}

Examples of GOOD statements:
- "Atoms share pairs of electrons to form covalent bonds"
- "Metal atoms lose electrons to become positive ions"
- "Electrons move around the outside of atoms"

Examples of BAD statements (too advanced):
- "Electronegativity differences determine bond polarity"
- "Delocalised electrons exhibit metallic characteristics"
- "Hybridisation affects molecular geometry"

Keep it simple for 14-15 year olds!"""

        try:
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a GCSE chemistry teacher for 14-15 year olds. Use ONLY simple GCSE-level language - NO university terms. Respond only with valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.7
            )
            
            content = response.choices[0].message.content
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
            return None
        except Exception as e:
            print(f"Error generating content: {e}")
            return None
    
    @staticmethod
    def generate_think_aloud_script(topic, example):
        """Generate comprehensive think-aloud teaching script"""
        prompt = f"""Create a 15-minute think-aloud teaching script for Year 10 GCSE chemistry students (ages 14-15).
Topic: {topic}
Example: {example}

CRITICAL REQUIREMENTS:
- Use ONLY GCSE-level terminology - NO university words
- Keep language simple and clear for 14-15 year olds
- Use everyday analogies they can relate to
- Avoid terms like: electronegativity, dipole, intermolecular forces (unless specifically teaching these)
- Use terms like: sharing electrons, full outer shell, strong/weak bonds, attracted to

Structure (follow I DO → WE DO → YOU DO pedagogy):

INTRO: (2-3 min)
[Engaging hook using everyday examples, clear learning goal, why it matters to them]

I DO PHASE:
STEP1: (2 min) [Teacher shows example step-by-step with [PAUSE] and _simple key terms_]
STEP2: (2 min) [Teacher explains WHY, include [I SAY YOU SAY] for key terms]
STEP3: (2 min) [Teacher completes example, [TURN AND TALK] - what did you notice?]

WE DO PHASE:
STEP4: (2 min) [Teacher guides class through similar example together]
STEP5: (2 min) [Practice another together, [MINI WHITEBOARD] - quick check]

YOU DO PHASE:
STEP6: (2 min) [[MINI WHITEBOARD] - students try independently, teacher circulates]
STEP7: (1 min) [[COLD CALL] ___ and ___, share your answers]

REVIEW: (2 min)
[Quick recap of key points, [COLD CALL] to check, link to next lesson]

Interactive elements:
- [I SAY YOU SAY]: "I say 'covalent bond', you say 'covalent bond'"
- [TURN AND TALK]: "Turn to your partner: what do you think happens next?"
- [MINI WHITEBOARD]: "On your boards, draw..." or "Write down..."
- [COLD CALL]: "___, what's the answer?" (only AFTER teaching the content)
- [THINKING]: [Hmm], [PAUSE], [Ah!], [points at board]
- _Underline_ key vocabulary for emphasis

Teaching techniques:
- Use memory tricks (e.g., "COValent = atoms Co-Operate and share")
- Address misconceptions directly: "Many students think... but actually..."
- Use analogies from daily life (sharing toys, tug-of-war, etc.)
- Break complex ideas into small, simple steps
- Repeat key terms 3-4 times throughout
- Ask simple yes/no or one-word answer questions

Make it sound like a real teacher talking to teenagers, not a textbook!"""

        try:
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a GCSE chemistry teacher for 14-15 year olds. Use ONLY simple, age-appropriate language. NO university-level terms. Think like you're explaining to teenagers, not university students. Keep it conversational and engaging."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.8
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Error generating script: {e}")
            return ""


class WordDocumentCreator:
    """Handles Word document creation and formatting"""
    
    @staticmethod
    def add_formatted_text(paragraph, text):
        """Add text with formatting markers and underlines"""
        if not text:
            return
        
        parts = text.split('_')
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                WordDocumentCreator._add_markers(paragraph, part)
            else:  # Underlined text
                run = paragraph.add_run(part)
                run.underline = True
                run.bold = True
    
    @staticmethod
    def _add_markers(paragraph, text):
        """Handle [MARKER] formatting in blue bold"""
        marker_parts = text.split('[')
        for j, part in enumerate(marker_parts):
            if j == 0:
                if part:
                    paragraph.add_run(part)
            else:
                if ']' in part:
                    marker, rest = part.split(']', 1)
                    run = paragraph.add_run(f"[{marker}]")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    if rest:
                        paragraph.add_run(rest)
    
    @staticmethod
    def create_document(topic, example, tkt_statements, tbat_statements, common_errors, script_text, output_file):
        """Create complete Word document"""
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Title
        title = doc.add_heading(f'Think-Aloud Script: {topic}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f'Example: {example}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True
        
        doc.add_paragraph()
        
        # Script sections
        doc.add_heading('Think-Aloud Script (15 minutes)', level=1)
        sections = WordDocumentCreator._parse_script(script_text)
        
        if 'intro' in sections:
            doc.add_heading('Introduction', level=2)
            WordDocumentCreator.add_formatted_text(doc.add_paragraph(), sections['intro'])
        
        for i in range(1, 8):
            step_key = f'step{i}'
            if step_key in sections:
                doc.add_heading(f'Step {i}', level=2)
                WordDocumentCreator.add_formatted_text(doc.add_paragraph(), sections[step_key])
        
        if 'review' in sections:
            doc.add_heading('Review', level=2)
            WordDocumentCreator.add_formatted_text(doc.add_paragraph(), sections['review'])
        
        doc.save(output_file)
        print(f"✓ Saved: {output_file}")
    
    @staticmethod
    def _parse_script(script_text):
        """Parse script into sections"""
        sections = {}
        current_section = None
        current_text = []
        
        for line in script_text.split('\n'):
            line = line.strip()
            if line.startswith('INTRO:'):
                current_section = 'intro'
                current_text = []
            elif line.startswith('STEP'):
                if current_section and current_text:
                    sections[current_section] = '\n'.join(current_text)
                current_section = line.split(':')[0].lower()
                current_text = []
            elif line.startswith('REVIEW:'):
                if current_section and current_text:
                    sections[current_section] = '\n'.join(current_text)
                current_section = 'review'
                current_text = []
            elif line:
                current_text.append(line)
        
        if current_section and current_text:
            sections[current_section] = '\n'.join(current_text)
        
        return sections


def generate_all_think_alouds():
    """Main function to generate all think-aloud documents"""
    doc_path = "C4.1 Unit Preparation Booklet.docx"
    excel_path = "C4.1 Unit Overview (4).xlsx"
    
    # Extract lesson data
    extractor = LessonContentExtractor()
    document_text = extractor.read_word_doc(doc_path)
    excel_objectives = extractor.read_excel_objectives(excel_path)
    lessons = extractor.extract_lessons(document_text)
    
    print(f"Found {len(lessons)} lessons\n")
    
    # Generate documents
    generator = AIContentGenerator()
    doc_creator = WordDocumentCreator()
    
    for lesson in lessons:
        lesson_num = lesson['lesson_number']
        print(f"Processing Lesson {lesson_num}: {lesson['title']}")
        
        # Get content from Excel or generate with AI
        excel_content = excel_objectives.get(lesson_num, {})
        
        if excel_content.get('know') or excel_content.get('do'):
            lesson_data = {
                'topic': lesson['title'],
                'example': 'H₂O',
                'tkt_statements': excel_content.get('know', []),
                'tbat_statements': excel_content.get('do', []),
                'common_errors': excel_content.get('common_errors', '')
            }
        else:
            print("  Generating content with AI...")
            lesson_data = generator.generate_lesson_metadata(lesson['title'])
            if not lesson_data:
                print("  Failed to generate content, skipping...")
                continue
        
        # Generate script
        print("  Generating script...")
        script = generator.generate_think_aloud_script(lesson_data['topic'], lesson_data['example'])
        
        # Create document
        safe_title = re.sub(r'[^\w\s-]', '', lesson['title'].lower()).replace(' ', '_')[:50]
        output_file = f"lesson_{lesson_num}_{safe_title}_think_aloud.docx"
        
        doc_creator.create_document(
            lesson_data['topic'],
            lesson_data['example'],
            lesson_data['tkt_statements'],
            lesson_data['tbat_statements'],
            lesson_data['common_errors'],
            script,
            output_file
        )
        
        time.sleep(1)  # Rate limiting
    
    print(f"\n✓ Complete! Generated {len(lessons)} documents")


if __name__ == "__main__":
    generate_all_think_alouds()