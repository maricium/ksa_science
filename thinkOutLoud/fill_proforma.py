"""Fill in lesson plan proforma using Anthropic API for expert teacher content"""

from docx import Document
from pathlib import Path
import sys
import os
import anthropic

# Add parent directory to path to import read_excel
sys.path.insert(0, str(Path(__file__).parent.parent))
from read_excel import get_lesson_data
from readnow.main import extract_slide6_objectives

# Load Anthropic API key
env_file = Path(__file__).parent.parent / '.env'
if env_file.exists():
    for line in env_file.read_text().splitlines():
        if line.startswith('ANTHROPIC_API_KEY='):
            os.environ['ANTHROPIC_API_KEY'] = line.split('=', 1)[1].strip('"').strip("'")

# Get API key - same method as lasttry.py
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY') or "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"
if not ANTHROPIC_API_KEY:
    print("❌ Error: ANTHROPIC_API_KEY not found")
    print("   Please add it to your .env file: ANTHROPIC_API_KEY=your_key_here")
    sys.exit(1)

client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


def generate_with_anthropic(prompt, max_tokens=2000):
    """Generate content using Anthropic Claude"""
    try:
        print("   🤖 Generating with Anthropic...")
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        print(f"   ⚠️  Anthropic API error: {e}")
        return None


def generate_misconceptions(lesson_title, know_content, objectives):
    """Generate common misconceptions using AI - expert teacher perspective"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience teaching Year 10-11 students (ages 14-16).

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

Learning objectives:
{objectives[:800]}

Generate a list of 5-7 common misconceptions that GCSE students typically have about this topic. For each misconception, provide:
1. The misconception (what students incorrectly think)
2. The reality (the correct understanding)
3. How to address it (specific teaching strategies you've found effective over 20 years)

Write in a practical, experienced teacher voice. Use specific examples from AQA GCSE Chemistry. Be concise but thorough.

Format as a numbered list with clear sections for each misconception."""
    
    result = generate_with_anthropic(prompt, max_tokens=1500)
    if result:
        return result
    return "Common misconceptions will be identified during lesson planning."


def generate_independent_task_script(lesson_title, know_content, do_content, objectives):
    """Generate a 5-minute think-aloud script using AI - expert teacher perspective"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience. Create a 5-minute "Completed Independent Task" think-aloud script.

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

What students need to do:
{do_content[:500]}

Learning objectives:
{objectives[:800]}

Create a script where you (the teacher) model completing a typical independent task for this lesson. The script should be EXACTLY 5 minutes when spoken.

Structure:
- STEP 1: Read and understand (30 seconds)
- STEP 2: Identify key information (1 minute)
- STEP 3: Work through the problem (2 minutes)
- STEP 4: Check and verify (1 minute)
- STEP 5: Reflect and extend (30 seconds)

Use these interactive elements throughout:
- [PAUSE] - for thinking pauses
- [THINKING] [Hmm], [Ah!] - internal thoughts
- [I SAY YOU SAY] - vocabulary practice
- [TURN AND TALK] - partner discussion
- [COLD CALL] - checking understanding
- [points at board] - visual references
- _Underline_ key terms for emphasis

Write in a natural, experienced teacher voice. Show your thinking process step-by-step. Include specific chemistry examples relevant to AQA GCSE. Make it practical and engaging for 14-16 year olds.

The script should model the exact thinking process students should use when working independently."""
    
    result = generate_with_anthropic(prompt, max_tokens=2000)
    if result:
        return result
    return f"5-minute independent task script for {lesson_title}"


def generate_i_do_script(lesson_title, know_content, do_content, objectives):
    """Generate I Do (teacher models) script with detailed steps - matching ideal proforma style"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience. Create a detailed "I Do" script where you model the learning for students.

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

What students need to do:
{do_content[:500]}

Learning objectives:
{objectives[:800]}

Create a detailed, practical script (approximately 8-10 minutes) where you (the teacher) model the key concept or skill. Write it like you're talking to your class - natural, flowing, with specific examples.

Include:
- What example(s) you'll show (be specific - e.g., "magnesium + copper sulfate")
- Step-by-step thinking process as you work through it
- Key questions you'll ask yourself aloud
- Vocabulary you'll emphasize
- What you'll point out on the board
- How you'll check understanding as you go

Write in a natural, conversational teacher voice. Include specific AQA GCSE Chemistry examples. Make it practical - another teacher should be able to follow it exactly.

Format: Write as flowing prose with clear sections, not numbered steps. Include specific examples, questions, and what to say at each stage."""
    
    result = generate_with_anthropic(prompt, max_tokens=2500)
    if result:
        return result
    return f"I Do script for {lesson_title}"


def generate_we_do_script(lesson_title, know_content, do_content, objectives):
    """Generate We Do (guided practice) script - matching ideal proforma style"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience. Create a detailed "We Do" script for guided practice.

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

What students need to do:
{do_content[:500]}

Learning objectives:
{objectives[:800]}

Create a detailed, practical script (approximately 6-8 minutes) where you guide the class through a similar example together.

Include:
- What example you'll work through together (be specific)
- Guided questions you'll ask: "Which metal is more reactive?", "What happens next?", etc.
- How students will contribute (mini whiteboards, calling out, etc.)
- What you'll write on the board
- How you'll check understanding at each stage

Write in a natural, conversational teacher voice. Include specific AQA GCSE Chemistry examples. Make it practical - list the guided questions you'll ask, what students will do, and what you'll say.

Format: Write as flowing prose with clear sections. Include specific guided questions in quotes, activities, and what to say."""
    
    result = generate_with_anthropic(prompt, max_tokens=2500)
    if result:
        return result
    return f"We Do script for {lesson_title}"


def generate_cfu_script(lesson_title, know_content, do_content, objectives):
    """Generate Check for Understanding script - matching ideal proforma style"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience. Create a detailed "Check for Understanding" script.

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

What students need to do:
{do_content[:500]}

Learning objectives:
{objectives[:800]}

Create a detailed, practical script (approximately 3-5 minutes) for checking student understanding.

Include:
- Specific questions to ask (write them in quotes, e.g., "Write the products of...", "State which substance...")
- How students will respond (cold call, mini whiteboards, etc.)
- What to look for in their answers
- How to address misconceptions immediately if they arise
- What to do if students are secure vs. if they need re-teaching

Write in a natural, conversational teacher voice. Include specific AQA GCSE Chemistry questions and examples. Make it practical - list the exact questions to ask and how to respond.

Format: Write as flowing prose. Include specific questions in quotes, checking strategies, and how to address misconceptions."""
    
    result = generate_with_anthropic(prompt, max_tokens=2000)
    if result:
        return result
    return f"Check for Understanding script for {lesson_title}"


def generate_you_do_script(lesson_title, know_content, do_content, objectives):
    """Generate You Do (independent practice) script - matching ideal proforma style"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience. Create a detailed "You Do" script for independent practice.

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

What students need to do:
{do_content[:500]}

Learning objectives:
{objectives[:800]}

Create a detailed, practical script (approximately 8-10 minutes) for independent student practice.

Include:
- What tasks students will complete (be specific - list the actual activities)
- Clear instructions: "You have X minutes to..."
- What you'll look for as you circulate
- Support prompts for struggling students
- How you'll review answers at the end
- What misconceptions to watch for

Write in a natural, conversational teacher voice. Include specific AQA GCSE Chemistry tasks and examples. Make it practical - list exactly what students will do, what to look for, and what to say.

Format: Write as flowing prose. Include specific tasks students will complete, what to look for while circulating, and how to review."""
    
    result = generate_with_anthropic(prompt, max_tokens=2500)
    if result:
        return result
    return f"You Do script for {lesson_title}"


def generate_intro_section(lesson_title, know_content, objectives):
    """Generate intro section using AI"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience.

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

Learning objectives:
{objectives[:800]}

Generate the "Intro" section for a lesson plan with these four parts:
1. Introduce: How to introduce this lesson (engaging hook, clear learning goal)
2. Sequence: How this lesson fits into the unit/curriculum (what comes before, what comes after)
3. Importance: Why this topic matters for GCSE Chemistry and real-world applications
4. Enthuse: How to get students excited about learning this topic

Write in a practical, experienced teacher voice. Be specific to AQA GCSE Chemistry. Keep each section concise (2-3 sentences)."""
    
    result = generate_with_anthropic(prompt, max_tokens=1000)
    if result:
        return result
    return f"Intro section for {lesson_title}"


def generate_tkt_tbat(lesson_title, know_content, do_content):
    """Generate TKT and TBAT statements using AI"""
    prompt = f"""You are an expert AQA GCSE Chemistry teacher with 20 years of experience.

Lesson: {lesson_title}

What students need to know:
{know_content[:800]}

What students need to do:
{do_content[:500]}

Generate:
1. TKT (They Know That) statements - 5-7 knowledge statements students should know by the end of the lesson
2. TBAT (They'll Be Able To) statements - 3-5 skill statements students should be able to do

Format as:
TKT:
• [statement 1]
• [statement 2]
...

TBAT:
• [statement 1]
• [statement 2]
...

Write in clear, GCSE-appropriate language. Be specific to AQA GCSE Chemistry. Use action verbs for TBAT statements (e.g., write, identify, explain, calculate)."""
    
    result = generate_with_anthropic(prompt, max_tokens=1500)
    if result:
        return result
    # Fallback
    know_points = [k.strip() for k in know_content.split('•') if k.strip()][:5]
    do_points = [d.strip() for d in do_content.split('•') if d.strip()][:3]
    tkt_text = '\n'.join([f"• {k}" if not k.startswith('•') else k for k in know_points])
    tbat_text = '\n'.join([f"• {d}" if not d.startswith('•') else d for d in do_points])
    return f"TKT:\n{tkt_text}\n\nTBAT:\n{tbat_text}"


def fill_lesson_plan_proforma(lesson_code, template_path, output_path):
    """Fill in lesson plan proforma with AI-generated content"""
    doc = Document(template_path)
    
    # Get lesson data
    lesson = None
    if lesson_code.startswith('C4.2'):
        import pandas as pd
        # Try relative to thinkOutLoud, then parent directory
        excel_file = Path("../Lesson Resources/C4.2 Extraction of Metals/Unit Guidance/C4.2 Unit Overview/C4.2 Unit Overview.xlsx")
        if not excel_file.exists():
            excel_file = Path("Lesson Resources/C4.2 Extraction of Metals/Unit Guidance/C4.2 Unit Overview/C4.2 Unit Overview.xlsx")
        if excel_file.exists():
            df = pd.read_excel(excel_file, sheet_name='C4.2 Extraction of Metals', header=3).dropna(how='all')
            for idx, row in df.iterrows():
                if len(row) > 1 and str(row.iloc[1]).strip() == lesson_code:
                    lesson = {
                        'code': lesson_code,
                        'title': str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else lesson_code,
                        'know': str(row.iloc[3]).strip() if len(row) > 3 and pd.notna(row.iloc[3]) else '',
                        'do': str(row.iloc[5]).strip() if len(row) > 5 and pd.notna(row.iloc[5]) else ''
                    }
                    break
    
    if not lesson:
        lesson = get_lesson_data(lesson_code)
    
    if not lesson:
        print(f"❌ Could not find lesson data for {lesson_code}")
        return False
    
    # Try to find Lesson Resources relative to thinkOutLoud directory
    lesson_resources = Path("../Lesson Resources")
    if not lesson_resources.exists():
        lesson_resources = Path("Lesson Resources")
    objectives = extract_slide6_objectives(lesson_code, str(lesson_resources), 6)
    
    print(f"📋 Filling proforma for {lesson_code}: {lesson.get('title')}")
    print(f"   Using Anthropic API to generate expert teacher content...\n")
    
    # Generate all content with AI
    misconceptions = generate_misconceptions(lesson.get('title', ''), lesson.get('know', ''), objectives or '')
    tkt_tbat = generate_tkt_tbat(lesson.get('title', ''), lesson.get('know', ''), lesson.get('do', ''))
    intro = generate_intro_section(lesson.get('title', ''), lesson.get('know', ''), objectives or '')
    independent_task = generate_independent_task_script(lesson.get('title', ''), lesson.get('know', ''), lesson.get('do', ''), objectives or '')
    i_do = generate_i_do_script(lesson.get('title', ''), lesson.get('know', ''), lesson.get('do', ''), objectives or '')
    we_do = generate_we_do_script(lesson.get('title', ''), lesson.get('know', ''), lesson.get('do', ''), objectives or '')
    cfu = generate_cfu_script(lesson.get('title', ''), lesson.get('know', ''), lesson.get('do', ''), objectives or '')
    you_do = generate_you_do_script(lesson.get('title', ''), lesson.get('know', ''), lesson.get('do', ''), objectives or '')
    
    # Fill Table 0 (header info)
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        if len(table0.rows) > 1 and len(table0.rows[1].cells) > 1:
            table0.rows[1].cells[1].text = lesson.get('title', '')
            print(f"   ✓ Filled lesson title")
        
        # Add misconceptions to Table 0, Row 2
        if len(table0.rows) > 2 and len(table0.rows[2].cells) > 1:
            current = table0.rows[2].cells[1].text.strip()
            table0.rows[2].cells[1].text = f"{current}\n\n{misconceptions}" if current and 'Big picture' in current else misconceptions
            print(f"   ✓ Filled Misconceptions")
        
        # Add independent task script to Row 3
        if len(table0.rows) > 3 and len(table0.rows[3].cells) > 1:
            table0.rows[3].cells[1].text = independent_task
            print(f"   ✓ Filled Completed independent task script")
    
    # Fill Table 1 (main lesson plan)
    if len(doc.tables) > 1:
        table1 = doc.tables[1]
        
        # Row 2: TKT & TBAT
        if len(table1.rows) > 2 and len(table1.rows[2].cells) > 1:
            table1.rows[2].cells[1].text = tkt_tbat
            print(f"   ✓ Filled TKT & TBAT")
        
        # Row 4: Do Now
        if len(table1.rows) > 4 and len(table1.rows[4].cells) > 2:
            do_now_text = f"\"You have 5 minutes to read the read now and answer the questions for {lesson_code}\""
            table1.rows[4].cells[2].text = do_now_text
            print(f"   ✓ Filled Do Now")
        
        # Row 6: Intro section
        if len(table1.rows) > 6 and len(table1.rows[6].cells) > 2:
            table1.rows[6].cells[2].text = intro
            print(f"   ✓ Filled Intro section")
        
        # Row 7: I Do
        if len(table1.rows) > 7 and len(table1.rows[7].cells) > 2:
            table1.rows[7].cells[2].text = i_do
            print(f"   ✓ Filled I Do script")
        
        # Row 8: We Do
        if len(table1.rows) > 8 and len(table1.rows[8].cells) > 2:
            table1.rows[8].cells[2].text = we_do
            print(f"   ✓ Filled We Do script")
        
        # Row 9: CfU (Check for Understanding)
        if len(table1.rows) > 9 and len(table1.rows[9].cells) > 2:
            table1.rows[9].cells[2].text = cfu
            print(f"   ✓ Filled CfU script")
        
        # Row 10: You Do
        if len(table1.rows) > 10 and len(table1.rows[10].cells) > 2:
            table1.rows[10].cells[2].text = you_do
            print(f"   ✓ Filled You Do script")
    
    doc.save(output_path)
    print(f"\n✅ Saved filled proforma to: {output_path}")
    return True


if __name__ == "__main__":
    # Paths relative to thinkOutLoud directory
    template = "lesson plan profoma GTR.docx"
    output = "C4.2.4_lesson_plan_proforma.docx"
    
    fill_lesson_plan_proforma("C4.2.4", template, output)
