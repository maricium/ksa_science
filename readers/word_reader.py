"""
Word Document Reader
Consolidated utilities for finding, reading, and creating Word documents

This module provides reusable functions for:
- Finding Word documents (Unit Preparation Booklets)
- Reading Word documents
- Extracting keyword definitions from Word documents
- Creating and saving ReadNow Word documents

Usage:
    from readers.word_reader import find_word_document, get_keyword_definition_from_word_doc

    # Find a Word document for a unit
    word_file = find_word_document("B3.2")

    # Get keyword definition from Word document
    definition = get_keyword_definition_from_word_doc("pyramid of biomass", "B3.2")
"""

from __future__ import annotations

import logging
import os
import re
from functools import lru_cache
from pathlib import Path

from docx import Document

logger = logging.getLogger("readers.word_reader")
import contextlib

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@lru_cache(maxsize=64)
def find_word_document(unit: str) -> Path | None:
    """
    Find Unit Preparation Booklet Word document for a given unit code.

    Args:
        unit: Unit code (e.g., "B3.2", "C4.1")

    Returns:
        Path to Word document or None if not found
    """
    # Find Unit Preparation Booklet Word document
    lesson_resources = Path("Lesson Resources")
    if not lesson_resources.exists():
        lesson_resources = Path("..") / "Lesson Resources"
    if not lesson_resources.exists():
        lesson_resources = Path("../..") / "Lesson Resources"

    # Look for Unit Preparation Booklet
    doc_files = list(lesson_resources.rglob(f"**/Unit Guidance/**/{unit}*Preparation*.docx"))
    if not doc_files:
        doc_files = list(lesson_resources.rglob(f"**/Unit Guidance/**/*{unit}*Preparation*.docx"))

    # Filter out temporary files
    doc_files = [f for f in doc_files if not f.name.startswith("~$")]

    if doc_files:
        return doc_files[0]

    return None


@lru_cache(maxsize=32)
def _read_word_document_cached(doc_path_str: str) -> Document | None:
    """
    Cached version of read_word_document (internal use).

    Args:
        doc_path_str: Path to Word document as string (for hashing)

    Returns:
        Document object or None if error
    """
    try:
        return Document(doc_path_str)
    except Exception as e:
        logger.warning(f"Error reading Word document: {e}")
        return None


def read_word_document(doc_path: Path) -> Document | None:
    """
    Read a Word document (with caching).

    Args:
        doc_path: Path to Word document

    Returns:
        Document object or None if error
    """
    return _read_word_document_cached(str(doc_path.absolute()))


def generate_keyword_variations(keyword: str) -> list[str]:
    """
    Generate variations of a keyword for searching (singular/plural, with/without "of", etc.).

    Args:
        keyword: The keyword to generate variations for

    Returns:
        List of keyword variations
    """
    keyword_lower = keyword.lower()
    variations = [
        keyword_lower,
        keyword_lower.replace(" of ", " "),
        keyword_lower + "s",  # plural
        keyword_lower.replace("s", ""),  # singular
    ]
    # Also try capitalized versions
    variations.extend([v.capitalize() for v in variations])
    return variations


def extract_definition_from_paragraph(text: str, keyword: str, matched_keyword: str) -> str | None:
    """
    Extract definition from a paragraph text using various patterns.

    Args:
        text: Paragraph text
        keyword: Original keyword
        matched_keyword: The keyword variation that was matched

    Returns:
        Definition string if found, None otherwise
    """
    text_lower = text.lower()

    # Pattern 1: "Keyword is..." or "Keyword: ..." or "Keyword means..."
    for kw_var in [matched_keyword, keyword]:
        definition_patterns = [
            rf"{re.escape(kw_var)}\s*[:\-–]\s*(.+?)(?:\.|$|;)",
            rf"{re.escape(kw_var)}\s+is\s+(.+?)(?:\.|$|;)",
            rf"{re.escape(kw_var)}\s+means\s+(.+?)(?:\.|$|;)",
            rf"{re.escape(kw_var)}\s+refers to\s+(.+?)(?:\.|$|;)",
            rf"{re.escape(kw_var)}\s+is called\s+(.+?)(?:\.|$|;)",
        ]

        for pattern in definition_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                definition = match.group(1).strip()
                if 10 < len(definition) < 300:  # Reasonable length
                    return definition

    # Pattern 2: Check if the paragraph itself is a definition (contains "is", "means", etc.)
    if any(word in text_lower for word in ["is", "means", "refers to", "called", "defined as"]):
        # Extract the part after the keyword
        parts = re.split(rf"{re.escape(keyword)}\s*[:\-–]?\s*", text, flags=re.IGNORECASE, maxsplit=1)
        if len(parts) > 1:
            definition = parts[1].strip()
            if 10 < len(definition) < 300:
                return definition

    return None


def extract_definition_from_following_paragraphs(
    doc: Document, start_idx: int, matched_keyword: str, keyword: str
) -> str | None:
    """
    Extract definition from paragraphs following a heading or keyword match.

    Args:
        doc: Document object
        start_idx: Starting paragraph index
        matched_keyword: The keyword variation that was matched
        keyword: Original keyword

    Returns:
        Definition string if found, None otherwise
    """
    # Pattern 3: If this looks like a heading, check next few paragraphs for definition
    current_text = doc.paragraphs[start_idx].text.strip()
    if len(current_text) < 100 and (current_text.isupper() or current_text.istitle()):
        # This might be a heading - check following paragraphs
        for j in range(1, 5):
            if start_idx + j < len(doc.paragraphs):
                next_text = doc.paragraphs[start_idx + j].text.strip()
                if (
                    next_text
                    and 20 < len(next_text) < 400
                    and (
                        re.match(rf"^{re.escape(matched_keyword)}\s+", next_text, re.IGNORECASE)
                        or re.match(r"^(is|means|refers to|called|defined as|are)", next_text, re.IGNORECASE)
                    )
                ):
                    # Extract definition part
                    if re.match(rf"^{re.escape(matched_keyword)}\s+", next_text, re.IGNORECASE):
                        # Remove keyword and get definition
                        definition = re.sub(
                            rf"^{re.escape(matched_keyword)}\s+", "", next_text, flags=re.IGNORECASE
                        ).strip()
                        if 10 < len(definition) < 300:
                            return definition
                    # Starts with definition word
                    elif 10 < len(next_text) < 300:
                        return next_text

    # Pattern 4: Check next few paragraphs for definition (general case)
    for j in range(1, 4):
        if (
            start_idx + j < len(doc.paragraphs)
            and (next_text := doc.paragraphs[start_idx + j].text.strip())
            and len(next_text) < 300
            and any(word in next_text.lower() for word in ["is", "means", "refers to"])
            and re.match(r"^(is|means|refers to|called|defined as)", next_text, re.IGNORECASE)
        ):
            return next_text

    return None


@lru_cache(maxsize=256)
def get_keyword_definition_from_word_doc(keyword: str, unit: str | None = None) -> str | None:
    """
    Get the definition for a specific keyword from the unit preparation Word document.

    Args:
        keyword: The keyword to find definition for
        unit: Unit code (e.g., "B3.2")

    Returns:
        Definition string if found, None otherwise
    """
    try:
        if not unit:
            return None

        # Find Word document
        doc_path = find_word_document(unit)
        if not doc_path:
            return None

        # Read Word document
        doc = read_word_document(doc_path)
        if not doc:
            return None

        # Generate keyword variations
        keyword_variations = generate_keyword_variations(keyword)

        # Search through all paragraphs for the keyword
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            text_lower = text.lower()

            # Check if this paragraph contains any variation of the keyword
            keyword_found = False
            matched_keyword = None
            for var in keyword_variations:
                if var in text_lower:
                    keyword_found = True
                    matched_keyword = var
                    break

            if keyword_found:
                # Try to extract definition from this paragraph
                definition = extract_definition_from_paragraph(text, keyword, matched_keyword)
                if definition:
                    return definition

                # Try to extract from following paragraphs
                definition = extract_definition_from_following_paragraphs(doc, i, matched_keyword, keyword)
                if definition:
                    return definition

        return None
    except Exception as e:
        logger.warning(f"Error reading Word document: {e}")
        return None


def get_subject_from_code(lesson_code: str) -> str:
    """
    Determine subject from lesson code prefix (C=Chemistry, P=Physics, B=Biology).

    Args:
        lesson_code: Lesson code (e.g., "C4.2.4", "P3.1.2", "B3.2.7")

    Returns:
        Subject name ("Chemistry", "Physics", "Biology", or "Unknown")
    """
    if lesson_code.startswith("C"):
        return "Chemistry"
    elif lesson_code.startswith("P"):
        return "Physics"
    elif lesson_code.startswith("B"):
        return "Biology"
    return "Unknown"


def save_readnow_docx(
    content: str,
    lesson_code: str,
    lesson_title: str,
    output_dir: str = "readnows",
    constants_module=None,
    get_subject_from_code_func=None,
    download_chemistry_image=None,
) -> bool:
    """
    Save ReadNow content to Word document organized by year and subject.

    Args:
        content: The ReadNow content text
        lesson_code: Lesson code (e.g., "C4.2.4")
        lesson_title: Lesson title
        output_dir: Output directory (default: "readnows")
        constants_module: Module containing constants (FONT_NAME, etc.)
        get_subject_from_code_func: Function to get subject from lesson code (defaults to get_subject_from_code)
        download_chemistry_image: Function to download chemistry image

    Returns:
        True if successful, False otherwise
    """
    try:
        # Import constants if not provided
        if constants_module is None:
            from constants import (
                FONT_NAME,
                FONT_SIZE_BODY,
                FONT_SIZE_BODY_SMALL,
                FONT_SIZE_HEADING,
                FONT_SIZE_TITLE,
                HEADING_COLOR,
                IMAGE_WIDTH,
                LINE_SPACING,
                MARGIN_LEFT_RIGHT,
                MARGIN_LEFT_RIGHT_HPA,
                MARGIN_TOP_BOTTOM,
                STUDENT_ATTAINMENT,
                STUDENT_YEAR,
                TAB_STOP_POSITION,
            )
        else:
            FONT_NAME = constants_module.FONT_NAME
            FONT_SIZE_TITLE = constants_module.FONT_SIZE_TITLE
            FONT_SIZE_HEADING = constants_module.FONT_SIZE_HEADING
            FONT_SIZE_BODY = constants_module.FONT_SIZE_BODY
            FONT_SIZE_BODY_SMALL = constants_module.FONT_SIZE_BODY_SMALL
            MARGIN_TOP_BOTTOM = constants_module.MARGIN_TOP_BOTTOM
            MARGIN_LEFT_RIGHT = constants_module.MARGIN_LEFT_RIGHT
            MARGIN_LEFT_RIGHT_HPA = constants_module.MARGIN_LEFT_RIGHT_HPA
            IMAGE_WIDTH = constants_module.IMAGE_WIDTH
            TAB_STOP_POSITION = constants_module.TAB_STOP_POSITION
            HEADING_COLOR = constants_module.HEADING_COLOR
            LINE_SPACING = constants_module.LINE_SPACING
            STUDENT_YEAR = constants_module.STUDENT_YEAR
            STUDENT_ATTAINMENT = constants_module.STUDENT_ATTAINMENT

        # Use provided function or default
        if get_subject_from_code_func is None:
            get_subject_from_code_func = get_subject_from_code

        year_folder = STUDENT_YEAR.replace(" ", "_")  # e.g., "Year_9"
        subject = get_subject_from_code_func(lesson_code)

        # Create directory structure: readnows/Year_9/Chemistry/
        final_output_dir = Path(output_dir) / year_folder / subject
        final_output_dir.mkdir(parents=True, exist_ok=True)

        is_hpa = STUDENT_ATTAINMENT.upper() == "HPA"
        suffix = "_HPA" if is_hpa else "_LPA"
        filepath = final_output_dir / f"{lesson_code}_ReadNow{suffix}.docx"

        def style_heading(heading):
            """Apply consistent styling to headings"""
            for run in heading.runs:
                run.font.size = FONT_SIZE_HEADING
                run.bold = True
                if not is_hpa:  # Only apply color for LPA
                    run.font.color.rgb = HEADING_COLOR

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = section.bottom_margin = MARGIN_TOP_BOTTOM
            section.left_margin = section.right_margin = MARGIN_LEFT_RIGHT_HPA if is_hpa else MARGIN_LEFT_RIGHT

        # Add title with dynamic date
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"Read Now: {lesson_title}")
        title_run.bold = True
        title_run.font.size = FONT_SIZE_TITLE
        title_run.font.name = FONT_NAME
        title_run.font.underline = WD_UNDERLINE.SINGLE  # Underline the title

        tab_stops = title_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(TAB_STOP_POSITION, WD_TAB_ALIGNMENT.RIGHT)
        title_para.add_run("\t")

        # Add "Date: " label first
        date_label_run = title_para.add_run("Date: ")
        date_label_run.bold = True
        date_label_run.font.name = FONT_NAME
        date_label_run.font.underline = WD_UNDERLINE.SINGLE  # Underline "Date: "

        # Add dynamic date field after the label
        run = title_para.add_run()
        run.font.underline = WD_UNDERLINE.SINGLE  # Underline the date field
        fldChar = OxmlElement("w:fldChar")  # creates a new element
        fldChar.set(qn("w:fldCharType"), "begin")  # sets attribute on element
        run._r.append(fldChar)  # appends element to run

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'DATE \\@ "d/M/yyyy" \\* MERGEFORMAT'  # Date format: 1/12/2025
        run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)

        # Add image
        image_path = None
        if download_chemistry_image:
            image_path = os.path.abspath(download_chemistry_image(lesson_title))
            if os.path.exists(image_path):
                try:
                    doc.add_paragraph()
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    img_para.add_run().add_picture(image_path, width=IMAGE_WIDTH)
                    doc.add_paragraph()
                except:
                    pass

        # Process content
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Skip if line is just the lesson title (to avoid duplicate title)
            # Check if the line matches the lesson title (with or without "Read Now:" prefix)
            title_variations = [
                lesson_title,
                f"Read Now: {lesson_title}",
                lesson_title.replace("(HT) ", "").replace("(FT) ", "").strip(),  # Remove HT/FT tags
            ]
            if any(line == var or (line.startswith(var) and len(line) <= len(var) + 20) for var in title_variations):
                continue
            elif line.upper() == "QUESTIONS":
                doc.add_paragraph()
                doc.add_paragraph()
                heading = doc.add_heading("Questions", level=1 if not is_hpa else 2)
                style_heading(heading)
            elif line.upper().startswith("MARK SCHEME"):
                doc.add_page_break()
                heading = doc.add_heading("MARK SCHEME", level=1 if not is_hpa else 2)
                style_heading(heading)
            else:
                para = doc.add_paragraph()
                if not is_hpa:  # Apply line spacing only for LPA
                    para.paragraph_format.line_spacing = LINE_SPACING

                # Process line to handle keywords marked with *keyword*
                parts = re.split(r"(\*[^*]+\*)", line)

                font_size = FONT_SIZE_BODY_SMALL if is_hpa else FONT_SIZE_BODY

                for part in parts:
                    if not part:
                        continue
                    run = para.add_run(part.strip("*"))
                    run.font.name = FONT_NAME
                    run.font.size = font_size
                    if part.startswith("*") and part.endswith("*") or line and line[0].isdigit() and "." in line[:3]:
                        run.bold = True

        doc.save(str(filepath))
        if image_path and "temp_image_" in image_path:
            with contextlib.suppress(BaseException):
                os.remove(image_path)
        logger.info(f"Saved to: {filepath.absolute()}")
        return True
    except Exception as e:
        logger.error(f"Error saving: {e}")
        import traceback

        traceback.print_exc()
        return False
